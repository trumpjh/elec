"""
제4회 문제 9 수동 추출 및 추가
"""
from docx import Document
import json

print("\n" + "="*70)
print("제4회 문제 9 추출 및 추가")
print("="*70)

# 1단계: 문제와 선택지 추출
print("\n【문제 및 선택지 추출】")
print("-" * 70)

doc = Document('2025년도 문제.docx')

# 문단 249-254: 문제 9와 선택지들
problem_text = doc.paragraphs[249].text.strip()  # 문제
print(f"문제: {problem_text}\n")

options = []
for idx in range(251, 255):
    option_text = doc.paragraphs[idx].text.strip()
    options.append(option_text)
    print(f"{option_text}")

# 2단계: 구조화된 데이터 생성
print("\n【데이터 정리】")
print("-" * 70)

# 문제 텍스트에서 답 추출
import re

# "? ③" 형태로 끝나므로 답 추출
match = re.search(r'\??\s*([①②③④])\s*$', problem_text)
if match:
    answer_symbol = match.group(1)
    answer_map = {'①': 0, '②': 1, '③': 2, '④': 3}
    answer_idx = answer_map[answer_symbol]
    print(f"답: {answer_symbol} (인덱스: {answer_idx})")
    
    # 문제 텍스트에서 답 기호 제거
    problem_clean = re.sub(r'\s*[①②③④]\s*$', '', problem_text)
    print(f"문제 (정리): {problem_clean[:60]}...")
else:
    print("❌ 답을 찾을 수 없습니다")
    answer_idx = 2  # 기본값

# 3단계: questions.json에 추가
print("\n【questions.json에 추가】")
print("-" * 70)

with open('questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 새로운 문제 객체 생성
new_problem = {
    'number': 9,
    'question': problem_clean,
    'options': options,
    'answer': answer_idx,
    'exam': '제4회',
    'explanation': '',  # 설명 추가 필요
    'category': '',  # 카테고리 추가 필요
    'image': 'images/problem_image_5.png'
}

# 제4회 문제 중에 맨 뒤에 추가 (또는 번호 순으로 정렬)
# 현재 제4회 문제들: 2, 3, 5, 6, 8, 10, 11, 12, 13, 14, 16, 17, 18, 19
# 9를 삽입할 위치 찾기

insert_pos = None
for i, q in enumerate(data['questions']):
    if q['exam'] == '제4회' and q['number'] == 10:
        insert_pos = i
        break

if insert_pos is not None:
    data['questions'].insert(insert_pos, new_problem)
    print(f"✓ 문제 추가됨 (위치: {insert_pos})")
else:
    data['questions'].append(new_problem)
    print(f"✓ 문제 추가됨 (마지막)")

# 저장
data['total'] += 1
with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print(f"✓ questions.json 업데이트 완료 (총 {data['total']}개)")

# 4단계: 확인
print("\n【확인】")
print("-" * 70)

with open('questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

exam4_problems = [q for q in data['questions'] if q['exam'] == '제4회']
print(f"\n제4회 문제 수: {len(exam4_problems)}")
print(f"문제 번호들: {sorted([q['number'] for q in exam4_problems])}")

for q in sorted(exam4_problems, key=lambda x: x['number']):
    image_info = f" [이미지: {q['image']}]" if q.get('image') else ""
    print(f"  문제 {q['number']}: {q['question'][:45]}...{image_info}")

print("\n" + "="*70)
