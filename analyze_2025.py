from docx import Document
import json
import re

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

questions = []
lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

print("=== 2025년도 문제.docx 분석 ===\n")
print("총 문단 수:", len(lines))
print("\n첫 50개 문단:")
for i, line in enumerate(lines[:50]):
    print(f"{i}: {line[:100]}")

print("\n\n=== 테이블 정보 ===")
for idx, table in enumerate(doc.tables):
    print(f"\n테이블 {idx + 1}: {len(table.rows)}개 행")
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  Row {row_idx}, Cell {cell_idx}: {text[:100]}")

print("\n\n=== 전체 텍스트 ===")
for i, line in enumerate(lines):
    print(f"{i}: {line}")
