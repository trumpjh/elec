from docx import Document
import sys

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

print(f"Total paragraphs: {len(doc.paragraphs)}\n")
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if text.strip():
        display = (text[:100] + '...') if len(text) > 100 else text
        print(f"{i:3d}: {display}")
