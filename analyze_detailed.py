"""
docx 파일을 다시 분석하여 문제와 설명의 정확한 매칭을 확인
"""

from docx import Document
import re
from collections import defaultdict

print("\n" + "="*70)
print("📊 정밀 분석: 문제와 설명 매칭")
print("="*70)

# 1. 문제 분석
print("\n【1단계】문제 추출 (문제.docx)")
print("-" * 70)

doc_questions = Document('2025년도 문제.docx')
paragraphs = [p.text.strip() for p in doc_questions.paragraphs if p.text.strip()]

problems_by_exam = defaultdict(list)
current_exam = None

for i, text in enumerate(paragraphs):
    # 회차 감지
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f'제{match.group(1)}회'
            print(f"\n→ {current_exam} 발견")
        continue
    
    # 문제 찾기
    match = re.match(r'^(\d+)\.\s+(.+)$', text)
    if match and current_exam:
        problem_num = int(match.group(1))
        problem_text = match.group(2)
        
        # 정답 심볼 확인
        answer_symbol = None
        for symbol in ['①', '②', '③', '④']:
            if problem_text.endswith(symbol):
                answer_symbol = symbol
                break
        
        if answer_symbol:
            problems_by_exam[current_exam].append({
                'number': problem_num,
                'text': problem_text[:-len(answer_symbol)].strip()[:50] + "...",
                'answer': answer_symbol
            })

# 회차별 문제 출력
for exam in sorted(problems_by_exam.keys()):
    problems = problems_by_exam[exam]
    print(f"\n{exam}: {len(problems)}개")
    for p in problems[:5]:  # 처음 5개만 표시
        print(f"  문제 {p['number']}: {p['text']}")
    if len(problems) > 5:
        print(f"  ... 외 {len(problems)-5}개")

# 2. 설명 분석
print("\n\n【2단계】설명 추출 (설명.docx)")
print("-" * 70)

doc_explanations = Document('2025년도 설명.docx')
tables = doc_explanations.tables

print(f"\n총 {len(tables)}개 표 발견\n")

explanations = []
for idx, table in enumerate(tables):
    if table.rows and table.rows[0].cells:
        cell_text = table.rows[0].cells[0].text.strip()
        
        # 설명의 첫 줄 추출 (단원명)
        first_line = cell_text.split('\n')[0] if cell_text else ''
        
        explanations.append({
            'table_index': idx,
            'category': first_line,
            'preview': cell_text[:60] + "..." if len(cell_text) > 60 else cell_text
        })

# 설명 미리보기
print("처음 10개 설명:")
for exp in explanations[:10]:
    print(f"  표 #{exp['table_index']:2d}: [{exp['category']:6s}] {exp['preview']}")

# 3. 문제 개수 비교
print("\n\n【3단계】문제와 설명 개수 비교")
print("-" * 70)

total_problems = sum(len(v) for v in problems_by_exam.values())
total_explanations = len(explanations)

print(f"\n문제 총 개수: {total_problems}개")
print(f"설명 총 개수: {total_explanations}개")
print(f"차이: {total_problems - total_explanations}개")

if total_problems > total_explanations:
    print(f"\n⚠️  설명이 {total_problems - total_explanations}개 부족합니다!")
    print("→ 문제와 설명을 정확히 매칭해야 합니다.")

# 4. 제1회 문제 상세 확인
print("\n\n【4단계】제1회 문제 상세 (특히 16번 주목)")
print("-" * 70)

exam1_problems = problems_by_exam['제1회']
for p in exam1_problems:
    marker = "👉" if p['number'] == 16 else "  "
    print(f"{marker} 문제 {p['number']:2d}: {p['text']}")
