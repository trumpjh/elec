"""
2025년도 문제.docx 및 설명.docx를 정밀하게 분석하는 스크립트
"""

from docx import Document
import re
from collections import defaultdict

class DocxAnalyzer:
    def __init__(self, docx_path):
        self.doc = Document(docx_path)
        self.paragraphs = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        self.tables = self.doc.tables
        self.analysis = {}
    
    def analyze_structure(self):
        """전체 구조 분석"""
        print("\n" + "="*100)
        print(f"파일: {self.doc}")
        print("="*100)
        
        print(f"\n【기본 정보】")
        print(f"  총 문단: {len(self.doc.paragraphs)}")
        print(f"  비어있지 않은 문단: {len(self.paragraphs)}")
        print(f"  표 개수: {len(self.tables)}")
        
        return self
    
    def analyze_rounds(self):
        """회차별 정보 분석"""
        print(f"\n【회차별 분석】")
        
        rounds = defaultdict(list)
        current_round = None
        
        for para in self.paragraphs:
            # 회차 찾기
            if '제' in para and '회' in para and '문제' in para:
                match = re.search(r'제(\d+)회', para)
                if match:
                    current_round = f"제{match.group(1)}회"
                    print(f"\n  >>> {current_round}")
                    continue
            
            # 문제 번호 찾기
            match = re.match(r'^(\d+)\.\s+', para)
            if match:
                num = int(match.group(1))
                rounds[current_round].append(num)
        
        # 통계
        print(f"\n【회차별 통계】")
        for round_name in sorted(rounds.keys()):
            problems = sorted(set(rounds[round_name]))
            print(f"  {round_name}: {len(problems)}개 - {problems}")
        
        self.analysis['rounds'] = dict(rounds)
        return self
    
    def analyze_problems(self):
        """문제 상세 분석"""
        print(f"\n【문제 상세 구조】")
        
        problems = {}
        i = 0
        while i < len(self.paragraphs):
            para = self.paragraphs[i]
            
            # 문제 찾기
            match = re.match(r'^(\d+)\.\s+(.+)$', para)
            if match:
                num = int(match.group(1))
                question_text = match.group(2)
                
                # 정답 추출
                answer = None
                for symbol_idx, symbol in enumerate(['①', '②', '③', '④']):
                    if symbol in question_text:
                        pos = question_text.rfind(symbol)
                        after = question_text[pos + 1:].strip()
                        if len(after) == 0:
                            answer = symbol_idx
                            break
                
                # 다음 문단에서 선택지 찾기
                options = []
                if i + 1 < len(self.paragraphs):
                    next_para = self.paragraphs[i + 1]
                    option_matches = re.findall(r'[①②③④]\s*([^①②③④]+?)(?=[①②③④]|$)', next_para)
                    options = [m.strip() for m in option_matches if m.strip()]
                
                if num not in problems:
                    problems[num] = {
                        'question': question_text,
                        'options': options[:4],
                        'answer': answer,
                        'has_options': len(options) > 0,
                        'has_answer': answer is not None
                    }
            
            i += 1
        
        print(f"  발견된 문제: {len(problems)}개")
        print(f"\n  처음 3개 문제:")
        for num in sorted(problems.keys())[:3]:
            p = problems[num]
            print(f"    문제 {num}:")
            print(f"      - 텍스트: {p['question'][:60]}...")
            print(f"      - 선택지: {len(p['options'])}개 (있음: {p['has_options']})")
            print(f"      - 정답: {['①', '②', '③', '④'][p['answer']] if p['answer'] is not None else 'None'} (있음: {p['has_answer']})")
        
        self.analysis['problems'] = problems
        return self
    
    def analyze_tables(self):
        """표 구조 분석"""
        print(f"\n【표 구조 분석】")
        print(f"  표 개수: {len(self.tables)}")
        
        if len(self.tables) > 0:
            print(f"\n  처음 3개 표:")
            for idx, table in enumerate(self.tables[:3]):
                print(f"\n    표 #{idx}:")
                print(f"      - 크기: {len(table.rows)}행 x {len(table.columns)}열")
                
                # 첫 셀 내용
                if table.rows and table.rows[0].cells:
                    cell_text = table.rows[0].cells[0].text.strip()
                    preview = cell_text[:100] if cell_text else "[비어있음]"
                    print(f"      - 첫 셀: {preview}...")
        
        self.analysis['tables'] = len(self.tables)
        return self
    
    def show_summary(self):
        """요약 표시"""
        print(f"\n【분석 요약】")
        print(f"  총 문제 발견: {len(self.analysis.get('problems', {}))}")
        print(f"  총 표: {self.analysis.get('tables', 0)}")
        
        rounds = self.analysis.get('rounds', {})
        if rounds:
            total = sum(len(set(v)) for v in rounds.values())
            print(f"  회차별 총 문제: {total}개")
        
        return self

# 분석 실행
print("\n" + "▶ "*50)
print("【2025년도 문제.docx 분석】")
print("▶ "*50)

analyzer1 = DocxAnalyzer('2025년도 문제.docx')
analyzer1.analyze_structure().analyze_rounds().analyze_problems().analyze_tables().show_summary()

print("\n" + "▶ "*50)
print("【2025년도 설명.docx 분석】")
print("▶ "*50)

analyzer2 = DocxAnalyzer('2025년도 설명.docx')
analyzer2.analyze_structure().analyze_rounds().analyze_problems().analyze_tables().show_summary()
