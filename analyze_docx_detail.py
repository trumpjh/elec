from docx import Document
import re

doc = Document('2025년도 문제.docx')

print("=" * 70)
print("2025년도 문제.docx 상세 분석")
print("=" * 70)

# 문제 찾기
problems = []
explanations = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 문제 패턴: "1. 문제? ④"
    problem_match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])\s*$', text)
    if problem_match:
        problems.append({
            'number': int(problem_match.group(1)),
            'text': problem_match.group(2),
            'answer': problem_match.group(3),
            'para_idx': i
        })
    
    # 설명 시작 패턴: "설명:"
    if text.startswith('설명:'):
        explanations.append({
            'para_idx': i,
            'text': text
        })

print(f"\n📝 문제 개수: {len(problems)}개")
print(f"   - 범위: {problems[0]['number'] if problems else 'N/A'} ~ {problems[-1]['number'] if problems else 'N/A'}")

if problems:
    print(f"\n   문제 번호 분포:")
    by_number = {}
    for p in problems:
        by_number[p['number']] = by_number.get(p['number'], 0) + 1
    
    for num in sorted(by_number.keys()):
        if by_number[num] > 1:
            print(f"   - 문제 {num}: {by_number[num]}개 (중복!)")

print(f"\n📚 설명 개수: {len(explanations)}개")

# 테이블 확인
print(f"\n📊 테이블 개수: {len(doc.tables)}개")

# 테이블 분석
table_explanations = 0
for table_idx, table in enumerate(doc.tables):
    print(f"\n   테이블 {table_idx + 1}:")
    print(f"   - 행: {len(table.rows)}")
    print(f"   - 열: {len(table.columns)}")
    
    # 테이블 내용 샘플
    if len(table.rows) > 0:
        first_row = table.rows[0]
        cells_text = [cell.text[:20] for cell in first_row.cells]
        print(f"   - 첫 행: {cells_text}")
        
        # 단원명- 패턴 찾기
        for row in table.rows:
            for cell in row.cells:
                if '단원명-' in cell.text or '설명:' in cell.text:
                    table_explanations += 1
                    break

print(f"\n   테이블에서 발견된 설명: {table_explanations}개")

print("\n" + "=" * 70)
print("✓ 분석 완료")
print("=" * 70)

# 마지막 문제 정보
if problems:
    last_p = problems[-1]
    print(f"\n마지막 문제 정보:")
    print(f"  번호: {last_p['number']}")
    print(f"  문제: {last_p['text'][:50]}...")
    print(f"  단락 위치: {last_p['para_idx']}")
