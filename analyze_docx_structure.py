from docx import Document

doc = Document('2025년도 문제.docx')

print("=" * 70)
print("DOCX 파일 구조 분석")
print("=" * 70)

# 모든 테이블 개수
print(f"\n총 테이블 개수: {len(doc.tables)}")

# 각 테이블의 구조 확인
for table_idx, table in enumerate(doc.tables):
    print(f"\n테이블 {table_idx}:")
    print(f"  행 개수: {len(table.rows)}")
    print(f"  열 개수: {len(table.columns)}")
    
    # 첫 3개 행 출력
    for row_idx in range(min(3, len(table.rows))):
        row_text = table.rows[row_idx].cells[0].text[:50]
        print(f"  행 {row_idx}: {row_text}")
    
    # 제4회 찾기
    for row in table.rows:
        if any('제4회' in cell.text for cell in row.cells):
            print(f"  ➜ 제4회 발견!")
            break

print("\n" + "=" * 70)
print("단락(Paragraph) 내용 검색")
print("=" * 70)

# 모든 단락에서 "제4회"와 "8번" 찾기
for para_idx, para in enumerate(doc.paragraphs):
    if '제4회' in para.text:
        print(f"\n단락 {para_idx}: 제4회 발견")
        print(f"  내용: {para.text[:100]}")
        
        # 다음 10개 단락 출력
        for i in range(1, 11):
            if para_idx + i < len(doc.paragraphs):
                next_para = doc.paragraphs[para_idx + i]
                print(f"  +{i}: {next_para.text[:80]}")
