"""
docx 파일의 구조를 분석하는 스크립트
"""

from docx import Document

def analyze_docx_structure(docx_path, max_items=30):
    """docx 파일의 구조를 상세히 분석"""
    doc = Document(docx_path)
    
    print(f"\n{'='*80}")
    print(f"파일: {docx_path}")
    print(f"{'='*80}\n")
    
    # 1. 문단 분석
    print("[문단 분석]")
    for i, para in enumerate(doc.paragraphs[:max_items]):
        if para.text.strip():
            print(f"{i}: {para.text[:100]}")
    
    print(f"\n총 문단 수: {len(doc.paragraphs)}")
    
    # 2. 표 분석
    print(f"\n[표 분석]")
    print(f"표 개수: {len(doc.tables)}")
    
    for t_idx, table in enumerate(doc.tables[:2]):
        print(f"\n표 #{t_idx}: {len(table.rows)} 행 x {len(table.columns)} 열")
        print(f"{'─'*80}")
        
        for r_idx, row in enumerate(table.rows[:15]):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"[행{r_idx}, 열{c_idx}]: {text[:80]}")

# 분석 실행
print("\n【2025년도 문제.docx 분석】")
analyze_docx_structure('2025년도 문제.docx', max_items=20)

print("\n\n【2025년도 설명.docx 분석】")
analyze_docx_structure('2025년도 설명.docx', max_items=20)
