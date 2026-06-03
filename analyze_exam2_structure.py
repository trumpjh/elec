from docx import Document
import re

doc = Document('2025년도 문제.docx')

# 제2회 찾기
print("="*70)
print("제2회 2번 문제 주변 상세 정보 확인")
print("="*70)

in_exam2 = False
found_q2 = False
context_start = None

for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if '제2회' in text and '회' in text:
        in_exam2 = True
        print(f"\n[{para_idx}] 제2회 시작")
        continue
    
    if in_exam2 and text.startswith('2.'):
        found_q2 = True
        context_start = max(0, para_idx - 2)
        break

if found_q2:
    # 문제 2번 주변 10줄 출력
    print("\n문제 2번 주변 문단들:\n")
    for i in range(context_start, min(context_start + 10, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text:
            # 문단 내 테이블 확인
            tables_in_para = 0
            
            print(f"[{i}] 문단 내용 ({len(text)}자):")
            print(f"    {text[:100]}...")
            print()

# 표 확인
print("\n" + "="*70)
print("표 정보 확인 (제2회 2번 문제와 관련)")
print("="*70)

for table_idx, table in enumerate(doc.tables):
    if table_idx < 15:  # 첫 15개 표만 확인
        try:
            cell_text = table.rows[0].cells[0].text.strip()[:100] if table.rows else "표 없음"
            print(f"[표 {table_idx}] {cell_text}...")
        except:
            pass

# 직접 검색
print("\n" + "="*70)
print("docx 원본 상태 (선택지 부분 분석)")
print("="*70)

paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
for i, text in enumerate(paragraphs):
    if '제2회' in text and '회' in text:
        # 제2회부터 제3회까지의 모든 문단 출력
        print(f"\n제2회 ~ 제3회 사이 모든 문단:")
        for j in range(i, min(i + 100, len(paragraphs))):
            if '제3회' in paragraphs[j]:
                break
            if j - i < 50:  # 최대 50줄
                print(f"[{j-i}] {paragraphs[j][:80]}")
        break
