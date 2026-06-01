from docx import Document
import json
import re

doc = Document(r'2025년도 문제.docx')

# 문제 추출
lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

problems = []
i = 0

while i < len(lines):
    line = lines[i]
    match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])', line)
    if match:
        q_num = int(match.group(1))
        question = match.group(2).strip()
        answer_char = match.group(3)
        answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
        
        options = []
        i += 1
        
        # 선택지 수집
        while i < len(lines):
            opt_line = lines[i]
            if re.match(r'^\d+\.\s+(.+)\?\s+([④③②①])', opt_line):
                break
            
            opt_match = re.match(r'^([①②③④])\s+(.*)', opt_line)
            if opt_match:
                options.append(opt_match.group(2).strip())
            elif re.search(r'①.*②.*③.*④', opt_line):
                parts = re.split(r'(?=[②③④])', opt_line)
                for part in parts:
                    if part:
                        opt_match = re.match(r'^([①②③④])\s+(.*)', part)
                        if opt_match:
                            options.append(opt_match.group(2).strip())
                i += 1
                break
            
            i += 1
        
        problems.append({
            'number': q_num,
            'question': question,
            'options': options[:4],
            'answer': answer_idx
        })
    else:
        i += 1

# 설명 추출
explanations = []
for table_idx, table in enumerate(doc.tables):
    text = table.rows[0].cells[0].text.strip()
    lines_text = text.split('\n')
    
    first_line = lines_text[0]
    category_match = re.search(r'단원명-(.+)', first_line)
    category = category_match.group(1) if category_match else '미분류'
    
    # 나머지는 설명
    explanation_lines = []
    for line in lines_text[1:]:
        line = line.strip()
        if line and line.startswith('-'):
            explanation_lines.append(line[1:].strip())
        elif line and not line.startswith('-'):
            explanation_lines.append(line)
    
    explanation = '\n'.join(explanation_lines)
    
    explanations.append({
        'index': table_idx,
        'category': category,
        'explanation': explanation
    })

# 정리 출력
print("=== 문제와 설명 매칭 ===")
print(f"문제 개수: {len(problems)}")
print(f"설명 개수: {len(explanations)}")
print()

# 문제 13이 있는지 확인
problem_nums = [p['number'] for p in problems]
print(f"문제 번호들: {problem_nums}")
print()

# 만약 설명이 한개 더 많으면, 빠진 문제 번호 찾기
if len(explanations) > len(problems):
    print(f"설명이 {len(explanations) - len(problems)}개 더 많습니다!")
    print()
    
    # 13이 빠져있는지 확인
    if 13 not in problem_nums:
        print("문제 13이 빠져있습니다!")
        print(f"설명 6 (인덱스 6): {explanations[6]['explanation'][:100]}")
        print("^ 이 설명이 문제 13의 설명일 것 같습니다.")
        print()

# 올바른 매칭
print("=== 올바른 매칭 ===")
for i, prob in enumerate(problems):
    expl_idx = i
    # 문제 13 설명이 있으면 스킵
    if 13 not in problem_nums and prob['number'] > 12:
        expl_idx = i + 1
    
    expl = explanations[expl_idx] if expl_idx < len(explanations) else None
    if expl:
        print(f"문제 {prob['number']:2d} <- 설명 {expl_idx:2d} ({expl['category']})")
    else:
        print(f"문제 {prob['number']:2d} <- 설명 없음!")
