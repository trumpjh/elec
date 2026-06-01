from docx import Document
import json
import re
from collections import defaultdict

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

questions_data = []
explanations_data = {}

# 설명 추출 (테이블)
table_idx = 0
for idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text.startswith("설명:"):
                table_idx += 1
                lines_text = text.split('\n')
                
                first_line = lines_text[0]
                category = ""
                if "단원명-" in first_line:
                    category = first_line.split("단원명-")[1].strip()
                
                explanation_lines = []
                for line in lines_text[1:]:
                    cleaned = line.strip()
                    if cleaned:
                        if cleaned.startswith("-"):
                            cleaned = cleaned[1:].strip()
                        explanation_lines.append(cleaned)
                
                explanation = "\n".join(explanation_lines)
                explanations_data[table_idx] = {
                    "category": category,
                    "explanation": explanation
                }

# 모든 단락 수집
paragraphs_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        paragraphs_text.append(text)

# 문제 추출
problem_idx = 1
i = 0
current_exam = "기출문제"

while i < len(paragraphs_text):
    line = paragraphs_text[i]
    
    # 회차 표시 감지
    exam_match = re.search(r'제(\d+)회', line)
    if exam_match:
        current_exam = f"제{exam_match.group(1)}회"
        i += 1
        continue
    
    # 문제 패턴
    match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])\s*$', line)
    
    if match:
        question_num = int(match.group(1))
        question_text = match.group(2).strip()
        answer_char = match.group(3)
        answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
        
        i += 1
        options = []
        
        # 선택지 찾기
        if i < len(paragraphs_text) and re.search(r'①.*②.*③.*④', paragraphs_text[i]):
            option_line = paragraphs_text[i]
            parts = re.split(r'(?=[②③④])', option_line)
            for part in parts:
                part = part.strip()
                if part and part[0] in '①②③④':
                    options.append(part[1:].strip())
            i += 1
        else:
            while i < len(paragraphs_text):
                option_line = paragraphs_text[i].strip()
                if re.match(r'^[①②③④]', option_line):
                    match_opt = re.match(r'^[①②③④]\s+(.+)', option_line)
                    if match_opt:
                        options.append(match_opt.group(1))
                        i += 1
                    else:
                        break
                else:
                    break
        
        if len(options) == 4:
            exp_data = explanations_data.get(problem_idx, {"category": "", "explanation": ""})
            
            questions_data.append({
                'number': question_num,
                'question': question_text,
                'options': options,
                'answer': answer_idx,
                'exam': current_exam
            })
            
            explanations_data[problem_idx] = {
                'number': question_num,
                'category': exp_data['category'],
                'explanation': exp_data['explanation']
            }
            
            problem_idx += 1
    else:
        i += 1

# 수동으로 누락된 4개 문제 추가
missing_problems = [
    {
        'number': 13,
        'question': '22.9[kV-Y] 가공전선의 굵기는 단면적이 몇 [mm²]이상이어야 하는가(단, 동선의 경우이다.)',
        'options': ['몰라', '몰라', '몰라', '몰라'],  # 선택지는 docx에서 추출 필요
        'answer': 0,  # ①
        'exam': '제1회'
    },
    {
        'number': 39,
        'question': '보극이 없는 직류기 운전 중 중성점의 위치가 변하지 않는 경우는',
        'options': ['과부하', '전부하', '중부하', '무부하'],
        'answer': 0,  # ①
        'exam': '제1회'
    },
    {
        'number': 49,
        'question': '전주 외등을 전주에 부착하는 경우 전주 외등은 하단으로부터 몇[m] 이상 높이에 시설하여야 하는가(단, 교통지장이 있는 경우이다.)',
        'options': ['몰라', '몰라', '몰라', '몰라'],  # 선택지는 docx에서 추출 필요
        'answer': 3,  # ④
        'exam': '제1회'
    },
    {
        'number': 2,
        'question': '그림과 같이 대전된 에보나이트 막대를 박검전기의 금속판에 닿지 않도록 가깝게 가져갔을 때 금박이 열렸다면 다음 중 옳은 것은(단, A는 원판, B는 박, C는 에보나이트 막대이다.)',
        'options': ['몰라', '몰라', '몰라', '몰라'],  # 선택지는 docx에서 추출 필요
        'answer': 2,  # ③
        'exam': '제2회'
    }
]

# 누락된 문제들을 questions_data에 추가
for missing in missing_problems:
    # 이미 존재하는지 확인
    if not any(q['number'] == missing['number'] and q['exam'] == missing['exam'] for q in questions_data):
        # 선택지와 함께 추가해야 함 (일단 추가는 하되, docx에서 선택지를 다시 추출해야 함)
        pass

# 회차별 집계
by_exam = defaultdict(list)
for q in questions_data:
    by_exam[q['exam']].append(q['number'])

print("=" * 80)
print("✓ 현재 추출 상태")
print("=" * 80)

for exam in sorted(by_exam.keys(), key=lambda x: int(re.search(r'\d+', x).group()) if re.search(r'\d', x) else 0):
    counts = len(by_exam[exam])
    print(f"  {exam}: {counts}개")

print(f"\n총 문제: {len(questions_data)}개")

print("\n누락된 4개 문제:")
for m in missing_problems:
    print(f"  {m['exam']} 문제 {m['number']}: 정답 {['①', '②', '③', '④'][m['answer']]}")
