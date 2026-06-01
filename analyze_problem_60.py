from docx import Document
import re

doc = Document('2025년도 문제.docx')

# 문제 60 주변 확인
print("문제 60 상세 분석:")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '60.' in text and '특수 장소' in text:
        print(f"\n[{i}] 문제: {text}")
        print(f"     길이: {len(text)}")
        
        # 다음 5개 단락
        for k in range(i+1, min(i+6, len(doc.paragraphs))):
            next_text = doc.paragraphs[k].text.strip()
            if next_text:
                print(f"[{k}] {next_text}")
                print(f"     길이: {len(next_text)}")
                # 특수 문자 감지
                if any(c in next_text for c in '①②③④'):
                    print(f"     => 보기 포함")
                if '㉠' in next_text or '㉡' in next_text:
                    print(f"     => 특수 기호 포함")
