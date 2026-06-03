from docx import Document

doc = Document('2025년도 설명.docx')

# 모든 문단을 인덱스와 함께 출력
print(f"총 {len(doc.paragraphs)} 개 문단\n")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:  # 비어있지 않은 문단만
        print(f"{i:3d}: {text}")
