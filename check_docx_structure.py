from docx import Document

doc = Document('2025년도 문제.docx')

print(f"총 테이블 수: {len(doc.tables)}")
print(f"총 단락 수: {len(doc.paragraphs)}")

# 첫 번째 테이블 구조 확인
if len(doc.tables) > 0:
    print("\n" + "=" * 80)
    print("첫 번째 테이블 샘플:")
    print("=" * 80)
    
    first_table = doc.tables[0]
    print(f"행 수: {len(first_table.rows)}, 열 수: {len(first_table.columns)}")
    
    # 처음 3행 출력
    for row_idx, row in enumerate(first_table.rows[:3]):
        print(f"\n행 {row_idx}:")
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text.strip()[:100]
            print(f"  열 {cell_idx}: {text}")

# 단락 샘플 확인
print("\n" + "=" * 80)
print("단락 샘플 (처음 50개):")
print("=" * 80)

for i, para in enumerate(doc.paragraphs[:50]):
    text = para.text.strip()
    if text:
        print(f"{i}: {text[:100]}")
