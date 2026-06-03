from docx import Document

doc = Document('2025년도 설명.docx')

print(f"총 문단 수: {len(doc.paragraphs)}")
print("\n첫 30개 문단:")
for i, para in enumerate(doc.paragraphs[:30]):
    text = para.text.strip()
    if text:
        print(f"{i}: {text[:80]}")
