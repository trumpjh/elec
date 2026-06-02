"""
문단 250 근처 내용 확인
"""
from docx import Document
import re

doc = Document('2025년도 문제.docx')

print("\n【문단 245-255 내용】")
print("-" * 70)

for idx in range(245, 256):
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        text = para.text[:60] if para.text else "[빈 문단]"
        
        # 이미지 확인
        has_image = any(run._element.drawing_lst for run in para.runs)
        image_marker = " [이미지]" if has_image else ""
        
        # 문제 번호 확인
        match = re.match(r'^(\d+)\.\s+', para.text)
        problem_marker = f" [문제 {match.group(1)}]" if match else ""
        
        print(f"{idx}: {text}...{image_marker}{problem_marker}")

print("\n【전체 문단 중 제4회 위치】")
print("-" * 70)

current_exam = None
exam_starts = {}

for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            exam_num = int(match.group(1))
            current_exam = f'제{exam_num}회'
            exam_starts[current_exam] = idx
            print(f"{idx}: {current_exam} 시작")

if '제4회' in exam_starts:
    print(f"\n제4회 시작 위치: {exam_starts['제4회']}")
    if '제5회' in exam_starts:
        print(f"제5회 시작 위치: {exam_starts.get('제5회', '없음')}")
    else:
        print(f"제4회 끝 위치: {len(doc.paragraphs) - 1} (문서 끝)")
