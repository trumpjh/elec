from docx import Document
import json

# DOCX 파일에서 제4회 8번 찾기
doc = Document('2025년도 문제.docx')

print("=" * 70)
print("DOCX 파일에서 제4회 8번 문제 찾기")
print("=" * 70)

found = False
for table_idx, table in enumerate(doc.tables):
    # 테이블의 첫 번째 셀이 제4회를 포함하는지 확인
    if table.rows:
        first_cell = table.cell(0, 0).text
        if '제4회' in first_cell or '4회' in first_cell:
            # 이 테이블에서 8번을 찾기
            for row in table.rows:
                row_text = ' '.join([cell.text for cell in row.cells])
                if '8' in row_text and '8번' in row_text:
                    print(f"\n테이블 {table_idx}: 제4회 8번 발견\n")
                    print("전체 행 내용:")
                    for i, cell in enumerate(row.cells):
                        print(f"  셀 {i}: {cell.text[:100]}")
                    found = True
                    break
            if found:
                break

if not found:
    print("\n제4회 테이블 찾아서 8번 행 확인:")
    for table_idx, table in enumerate(doc.tables):
        if table.rows:
            first_cell = table.cell(0, 0).text
            if '제4회' in first_cell or '4회' in first_cell:
                print(f"\n테이블 {table_idx} (제4회)의 모든 행:")
                for row_idx, row in enumerate(table.rows[:20]):  # 처음 20개 행만
                    row_text = ' '.join([cell.text[:50] for cell in row.cells])
                    print(f"  행 {row_idx}: {row_text}")
