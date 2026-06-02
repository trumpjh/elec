from docx import Document
import json

doc = Document('2025년도 문제.docx')

print("=" * 80)
print("제4회 문제 번호별 위치")
print("=" * 80)

# 제4회 시작 찾기
start_idx = None
for i, para in enumerate(doc.paragraphs):
    if '제4회' in para.text and '2025' in para.text:
        start_idx = i
        print(f"\n제4회 시작: 단락 {i}")
        break

if start_idx:
    # 제4회의 모든 내용을 한줄씩 출력
    print("\n제4회 전체 문제 내용:")
    print("-" * 80)
    
    for i in range(start_idx, min(start_idx + 100, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        # 다음 회차 발견 시 종료
        if i > start_idx and '제' in text and '회' in text and text.startswith('2025'):
            break
        
        if text:
            print(f"{i}: {text}")

# JSON 파일에서 제4회 10, 14, 19번의 현재 선택지 개수 확인
print("\n" + "=" * 80)
print("questions.json 현황")
print("=" * 80)

with open('questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

for q in data['questions']:
    if q['exam'] == '제4회' and q['number'] in [10, 14, 19]:
        print(f"\n제4회 {q['number']}번:")
        print(f"  선택지 개수: {len(q['options'])}")
        print(f"  문제: {q['question'][:50]}...")
        print(f"  보기: {q['options']}")
