from docx import Document
import json

doc = Document('2025년도 문제.docx')

print("=" * 80)
print("DOCX 파일의 테이블 구조 확인")
print("=" * 80)

# 모든 테이블 확인
print(f"\n총 테이블 개수: {len(doc.tables)}")

for table_idx, table in enumerate(doc.tables):
    print(f"\n【테이블 {table_idx}】")
    print(f"  행 개수: {len(table.rows)}")
    print(f"  열 개수: {len(table.columns)}")
    
    # 테이블 내용 출력
    for row_idx, row in enumerate(table.rows):
        print(f"  행 {row_idx}:")
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if cell_text:
                # 긴 텍스트는 처음 80자만 출력
                preview = cell_text[:80] + "..." if len(cell_text) > 80 else cell_text
                print(f"    [셀 {col_idx}] {preview}")

print("\n" + "=" * 80)
print("JSON 파일에서 해당 문제 확인")
print("=" * 80)

with open('questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

problems_to_check = [
    ('제1회', 60),
    ('제3회', 55),
    ('제3회', 58)
]

for exam, num in problems_to_check:
    for q in data['questions']:
        if q['exam'] == exam and q['number'] == num:
            print(f"\n【{exam} {num}번】")
            print(f"문제: {q['question'][:100]}...")
            print(f"선택지: {q['options']}")
            break
