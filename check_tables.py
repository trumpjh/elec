from docx import Document

doc = Document('2025년도 설명.docx')

# 모든 테이블을 자세히 확인
print("=== 모든 테이블 내용 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n테이블 {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            print(f"  [{r_idx},{c_idx}]: {text[:200]}")
            
# 44번과 59번이 테이블에 있는지 확인
print("\n\n=== 44번 검색 (테이블) ===")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if '44' in cell.text or '44.' in cell.text:
                print(f"테이블 {t_idx}, Row {r_idx}, Cell {c_idx}:")
                print(f"  {cell.text[:500]}")

print("\n=== 59번 검색 (테이블) ===")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if '59' in cell.text or '59.' in cell.text:
                print(f"테이블 {t_idx}, Row {r_idx}, Cell {c_idx}:")
                print(f"  {cell.text[:500]}")
