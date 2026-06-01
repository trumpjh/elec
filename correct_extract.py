from docx import Document
import json
import re

doc = Document(r'2025년도 문제.docx')

# 문제 추출
questions_list = []
lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

i = 0
problem_data = []

while i < len(lines):
    line = lines[i]
    
    # 문제 시작 패턴 찾기
    match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])', line)
    if match:
        q_num = int(match.group(1))
        question = match.group(2).strip()
        answer_char = match.group(3)
        answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
        
        options = []
        i += 1
        
        # 선택지 수집
        while i < len(lines):
            opt_line = lines[i]
            
            # 다음 문제 시작 패턴이면 멈춤
            if re.match(r'^\d+\.\s+(.+)\?\s+([④③②①])', opt_line):
                break
            
            # 선택지 라인
            opt_match = re.match(r'^([①②③④])\s+(.*)', opt_line)
            if opt_match:
                options.append(opt_match.group(2).strip())
                i += 1
            # 한 라인에 여러 선택지
            elif re.search(r'①.*②.*③.*④', opt_line):
                parts = re.split(r'(?=[②③④])', opt_line)
                for part in parts:
                    if part:
                        opt_match = re.match(r'^([①②③④])\s+(.*)', part)
                        if opt_match:
                            options.append(opt_match.group(2).strip())
                i += 1
                break
            else:
                i += 1
                break
        
        problem_data.append({
            'number': q_num,
            'question': question,
            'options': options[:4],
            'answer': answer_idx
        })
    else:
        i += 1

# 설명 추출
explanations = {}
for table_idx, table in enumerate(doc.tables):
    text = table.rows[0].cells[0].text.strip()
    lines = text.split('\n')
    
    # 첫 번째 줄에서 카테고리 추출
    first_line = lines[0]
    category_match = re.search(r'단원명-(.+)', first_line)
    category = category_match.group(1) if category_match else '미분류'
    
    # 나머지 부분은 설명
    explanation_lines = [l.strip() for l in lines[1:] if l.strip() and l.strip().startswith('-')]
    explanation = '\n'.join([l[2:] if l.startswith('- ') else l for l in explanation_lines])
    
    explanations[table_idx] = {
        'category': category,
        'explanation': explanation
    }

print("=== 문제와 설명 매칭 ===")
print(f"총 문제: {len(problem_data)}")
print(f"총 설명: {len(explanations)}")
print()

for idx, (prob, expl_data) in enumerate(zip(problem_data, explanations.values())):
    print(f"문제 {idx}: 번호={prob['number']}, 카테고리={expl_data['category']}")
    print(f"  질문: {prob['question'][:50]}...")
    print(f"  설명: {expl_data['explanation'][:80]}...")
    print()

# JSON 생성
questions_json = []
for prob, expl_data in zip(problem_data, explanations.values()):
    questions_json.append({
        'number': prob['number'],
        'question': prob['question'],
        'options': prob['options'],
        'answer': prob['answer'],
        'category': expl_data['category'],
        'explanation': expl_data['explanation']
    })

# JSON 저장
with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump({'questions': questions_json}, f, ensure_ascii=False, indent=2)

print("\n=== JSON 저장 완료 ===")
print(f"저장된 문제: {len(questions_json)}개")
