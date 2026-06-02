"""
이미지를 정확하게 매핑 (각 이미지를 구분해서 처리)
"""

from docx import Document
import re
import json
import os

print("\n" + "="*70)
print("🖼️  이미지 정확한 매핑")
print("="*70)

doc = Document('2025년도 문제.docx')

# 1단계: 이미지가 있는 문단과 문제 번호 매핑
print("\n【이미지 있는 문단과 문제 매칭】")
print("-" * 70)

image_problem_map = {}  # {para_idx: (exam, problem_num)}
current_exam = None

for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 회차 감지
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f'제{match.group(1)}회'
    
    # 이미지가 있는 문단 확인
    has_image = any(run._element.drawing_lst for run in para.runs)
    
    if has_image and current_exam:
        # 이 이미지가 어느 문제에 속하는지 찾기
        for i in range(para_idx - 1, -1, -1):
            prev_para = doc.paragraphs[i]
            prev_text = prev_para.text.strip()
            
            match = re.match(r'^(\d+)\.\s+(.+)$', prev_text)
            if match and prev_text:
                problem_num = int(match.group(1))
                image_problem_map[para_idx] = (current_exam, problem_num)
                print(f"  문단 {para_idx}: {current_exam} 문제 {problem_num} → 이미지 있음")
                break

# 2단계: 순서대로 이미지에 번호 부여
print("\n【이미지별 파일명 결정】")
print("-" * 70)

sorted_image_paras = sorted(image_problem_map.keys())
image_filenames = {}  # {para_idx: (exam, problem_num, filename)}

images_dir = 'images'

# 이미지 파일들 확인
image_files = sorted([f for f in os.listdir(images_dir) if f.startswith('problem_image_')])
print(f"저장된 이미지 파일: {image_files}")

for idx, para_idx in enumerate(sorted_image_paras):
    if idx < len(image_files):
        exam, problem_num = image_problem_map[para_idx]
        filename = image_files[idx]
        image_filenames[para_idx] = (exam, problem_num, filename)
        print(f"  문단 {para_idx}: {exam} 문제 {problem_num} ← {filename}")

# 3단계: questions.json 업데이트
print("\n【questions.json 업데이트】")
print("-" * 70)

with open('questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 먼저 모든 문제의 image 필드 초기화
for question in data['questions']:
    if 'image' not in question:
        question['image'] = None

# 이미지 정보 추가
for para_idx, (exam, problem_num, filename) in image_filenames.items():
    for question in data['questions']:
        if question['exam'] == exam and question['number'] == problem_num:
            question['image'] = f"images/{filename}"
            print(f"  {exam} 문제 {problem_num}: {filename} 연결")
            break

# 저장
with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print(f"\n✓ questions.json 업데이트 완료")

# 4단계: 최종 확인
print("\n【이미지 포함 문제】")
print("-" * 70)

problems_with_images = [q for q in data['questions'] if q.get('image')]
print(f"\n총 {len(problems_with_images)}개 문제에 이미지 포함:\n")

for q in problems_with_images:
    print(f"  {q['exam']} 문제 {q['number']}: {q['image']}")

print("\n" + "="*70)
print("✅ 모든 작업 완료!")
print("="*70)
