from docx import Document
import json
import re

doc = Document(r'2025년도 문제.docx')

# 문제 추출
lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

problems = []
i = 0

while i < len(lines):
    line = lines[i]
    match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])', line)
    if match:
        q_num = int(match.group(1))
        question = match.group(2).strip()
        answer_char = match.group(3)
        answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
        
        options = []
        i += 1
        
        # 선택지 수집
        while i < len(lines):
            opt_line = lines[i]
            if re.match(r'^\d+\.\s+(.+)\?\s+([④③②①])', opt_line):
                break
            
            opt_match = re.match(r'^([①②③④])\s+(.*)', opt_line)
            if opt_match:
                options.append(opt_match.group(2).strip())
            elif re.search(r'①.*②.*③.*④', opt_line):
                parts = re.split(r'(?=[②③④])', opt_line)
                for part in parts:
                    if part:
                        opt_match = re.match(r'^([①②③④])\s+(.*)', part)
                        if opt_match:
                            options.append(opt_match.group(2).strip())
                i += 1
                break
            
            i += 1
        
        problems.append({
            'number': q_num,
            'question': question,
            'options': options[:4],
            'answer': answer_idx
        })
    else:
        i += 1

# 설명 추출
explanations = []
for table_idx, table in enumerate(doc.tables):
    text = table.rows[0].cells[0].text.strip()
    lines_text = text.split('\n')
    
    first_line = lines_text[0]
    category_match = re.search(r'단원명-(.+)', first_line)
    category = category_match.group(1) if category_match else '미분류'
    
    # 나머지는 설명
    explanation_lines = []
    for line in lines_text[1:]:
        line = line.strip()
        if line and line.startswith('-'):
            explanation_lines.append(line[1:].strip())
        elif line and not line.startswith('-'):
            explanation_lines.append(line)
    
    explanation = '\n'.join(explanation_lines).strip()
    
    explanations.append({
        'category': category,
        'explanation': explanation
    })

# 올바른 매칭으로 JSON 생성
questions_json = []

for i, problem in enumerate(problems):
    prob_num = problem['number']
    
    # 문제 13이 빠져있으므로, 문제 14부터는 설명 인덱스가 한 칸씩 뒤로 밀림
    if prob_num > 12:
        expl_idx = i + 1  # 설명 6을 스킵하고 설명 7부터
    else:
        expl_idx = i
    
    expl = explanations[expl_idx]
    
    questions_json.append({
        'number': prob_num,
        'question': problem['question'],
        'options': problem['options'],
        'answer': problem['answer'],
        'category': expl['category'],
        'explanation': expl['explanation']
    })

# JSON 저장
with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump({'questions': questions_json}, f, ensure_ascii=False, indent=2)

print("=== 올바른 questions.json 생성 완료 ===")
print(f"저장된 문제: {len(questions_json)}개")
print()

# 확인 출력
for q in questions_json:
    print(f"문제 {q['number']:2d} ({q['category']:6s}): {q['question'][:50]}...")
