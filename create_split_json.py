from docx import Document
import json
import re

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

# 문제와 설명을 분리해서 저장
questions_data = []
explanations_data = {}

lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

# 설명 추출 (테이블에서)
table_idx = 0
for idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text.startswith("설명:"):
                table_idx += 1
                lines_text = text.split('\n')
                
                # 첫 번째 라인에서 단원명 추출
                first_line = lines_text[0]
                category = ""
                if "단원명-" in first_line:
                    category = first_line.split("단원명-")[1].strip()
                
                # 설명 추출
                explanation_lines = []
                for line in lines_text[1:]:
                    cleaned = line.strip()
                    if cleaned:
                        if cleaned.startswith("-"):
                            cleaned = cleaned[1:].strip()
                        explanation_lines.append(cleaned)
                
                explanation = "\n".join(explanation_lines)
                explanations_data[table_idx] = {
                    "category": category,
                    "explanation": explanation
                }

# 문제 추출 (문단에서)
problem_idx = 1
i = 0
current_exam = "기출문제"

while i < len(lines):
    line = lines[i]
    
    # 회차 제목 감지 - "제1회", "제2회", "제3회" 정확히 추출
    exam_match = re.search(r'제(\d+)회', line)
    if exam_match:
        current_exam = f"제{exam_match.group(1)}회"
        i += 1
        continue
    
    # 문제 찾기 (번호. 문제 형식)
    # 패턴 1: "번호. 문제? ④" (정규)
    match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])\s*$', line)
    # 패턴 2: "번호. 문제?" (정답 없음 - 다음 줄에서 찾기)
    match_no_answer = re.match(r'^(\d+)\.\s+(.+)\?\s*$', line) if not match else None
    
    if match:
        question_num = int(match.group(1))
        question_text = match.group(2).strip()
        answer_char = match.group(3)
        answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
        
        # 선택지 찾기
        i += 1
        options = []
        
        # 한 라인에 모든 선택지가 있는 경우
        if i < len(lines) and re.search(r'①.*②.*③.*④', lines[i]):
            option_line = lines[i]
            parts = re.split(r'(?=[②③④])', option_line)
            for part in parts:
                part = part.strip()
                if part and part[0] in '①②③④':
                    options.append(part[1:].strip())
            i += 1
        
        # 선택지가 여러 줄에 걸쳐 있는 경우
        else:
            while i < len(lines):
                option_line = lines[i].strip()
                if re.match(r'^[①②③④]', option_line):
                    match_opt = re.match(r'^[①②③④]\s+(.+)', option_line)
                    if match_opt:
                        options.append(match_opt.group(1))
                        i += 1
                    else:
                        break
                else:
                    break
        
        if len(options) == 4:
            # 설명 찾기
            exp_data = explanations_data.get(problem_idx, {"category": "", "explanation": ""})
            
            questions_data.append({
                'number': question_num,
                'question': question_text,
                'options': options,
                'answer': answer_idx,
                'exam': current_exam
            })
            
            # 설명은 따로 저장
            explanations_data[problem_idx] = {
                'number': question_num,
                'category': exp_data['category'],
                'explanation': exp_data['explanation']
            }
            
            problem_idx += 1
        else:
            i += 1
    else:
        i += 1

# questions.json 저장 (문제와 정답만)
questions_output = {
    'total': len(questions_data),
    'questions': questions_data
}

with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(questions_output, f, ensure_ascii=False, indent=2)

# example.json 저장 (설명)
# 문제 번호로 매핑
example_data = []
for idx in sorted(explanations_data.keys()):
    if explanations_data[idx]['explanation']:  # 설명이 있는 것만
        example_data.append({
            'problem_number': explanations_data[idx].get('number', idx),
            'category': explanations_data[idx]['category'],
            'explanation': explanations_data[idx]['explanation']
        })

example_output = {
    'total': len(example_data),
    'examples': example_data
}

with open('example.json', 'w', encoding='utf-8') as f:
    json.dump(example_output, f, ensure_ascii=False, indent=2)

print("=" * 80)
print("✓ 파일 생성 완료!")
print("=" * 80)
print(f"\n✓ questions.json: {len(questions_data)}개 문제")
print(f"✓ example.json: {len(example_data)}개 설명")
print("\n=== questions.json 구조 ===")
print(json.dumps(questions_data[0], ensure_ascii=False, indent=2))
print("\n=== example.json 구조 ===")
if example_data:
    print(json.dumps(example_data[0], ensure_ascii=False, indent=2))
