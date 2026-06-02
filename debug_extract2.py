"""
docx 파일에서 모든 문제를 debug 모드로 추출
"""

from docx import Document
import re

docx_path = '2025년도 문제.docx'
doc = Document(docx_path)

count = 0
current_exam = None

# 모든 문단 순회
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 회차 찾기
    if '제' in text and '회' in text and '문제' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f"제{match.group(1)}회"
            print(f"\n>>> {current_exam}를 찾았습니다")
    
    # 문제 번호 찾기: check_rounds.py와 동일한 정규식 사용
    match = re.match(r'^(\d+)\.\s+(.*)', text)
    if match:
        num = int(match.group(1))
        rest = match.group(2)
        count += 1
        
        print(f"문제 {num}: {rest[:60]}...")
        
        # 다음 문단에서 선택지 찾기
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i + 1].text.strip()
            if next_text and (next_text[0] in ['①', '②', '③', '④']):
                print(f"  선택지: {next_text[:60]}...")

print(f"\n\n총 발견된 문제: {count}")
