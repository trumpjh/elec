from docx import Document

doc = Document('2025년도 문제.docx')

# 문제 4 주변 단락 출력 (보기가 없는 것으로 탐지된 문제)
print("=" * 80)
print("문제 4 주변 텍스트 (인덱스 10-20):")
print("=" * 80)

for i in range(10, min(20, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if text:
        print(f"{i}: {text}")

# 문제 번호가 있는데 다음에 옵션이 없는 경우 찾기
print("\n" + "=" * 80)
print("보기가 없는 문제들 찾기:")
print("=" * 80)

import re
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    match = re.match(r'^(\d+)\.\s+(.+)\?\s*([①②③④])\s*$', text)
    
    if match:
        problem_num = int(match.group(1))
        
        # 다음 단락 확인
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i + 1].text.strip()
            
            # 다음이 옵션 라인이 아닌 경우
            if not re.search(r'^①.*②.*③.*④', next_text):
                # 그 다음 단락 확인
                if i + 2 < len(doc.paragraphs):
                    next_next = doc.paragraphs[i + 2].text.strip()
                    if re.search(r'^①.*②.*③.*④', next_next):
                        print(f"\n문제 {problem_num}:")
                        print(f"  [{i}] 질문: {text[:60]}")
                        print(f"  [{i+1}] (중간): {next_text[:80]}")
                        print(f"  [{i+2}] 보기: {next_next[:80]}")
                    else:
                        print(f"\n문제 {problem_num}: 보기 못 찾음")
                        print(f"  [{i}] 질문: {text[:60]}")
                        print(f"  [{i+1}] 다음: {next_text[:80]}")
                        if i + 2 < len(doc.paragraphs):
                            print(f"  [{i+2}] 그다음: {doc.paragraphs[i + 2].text.strip()[:80]}")
