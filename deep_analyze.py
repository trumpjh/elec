from docx import Document

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

# 테이블 확인
print(f"Tables: {len(doc.tables)}")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}: {len(table.rows)} rows, {len(table.columns)} columns")
    for row_idx, row in enumerate(table.rows[:5]):
        for col_idx, cell in enumerate(row.cells):
            print(f"  [{row_idx},{col_idx}]: {cell.text[:50]}")

# 모든 문단 출력
print("\n\nAll paragraphs:")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{i}: {para.text[:80]}")
