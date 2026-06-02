"""
docx 파일에서 정확하게 모든 문제를 추출하는 상세 분석
"""

from docx import Document
import re

def detailed_analysis(docx_path):
    """상세 파일 분석"""
    doc = Document(docx_path)
    
    print(f"\n파일: {docx_path}")
    print(f"{'='*80}")
    
    # 문단 분석
    print("\n【문단별 분석 (처음 100개)】")
    for i, para in enumerate(doc.paragraphs[:100]):
        text = para.text.strip()
        if text:
            # 회차 표시
            if '회' in text and '문제' in text:
                print(f"\n>>> {text}")
            # 문제 번호
            elif re.match(r'^\d+\.\s+', text):
                print(f"[{i}] {text[:70]}")
    
    # 표 상세 분석
    print(f"\n【표 상세 분석】")
    for t_idx, table in enumerate(doc.tables[:5]):
        print(f"\n표 #{t_idx}: {len(table.rows)}행 x {len(table.columns)}열")
        for r_idx, row in enumerate(table.rows[:3]):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"  [{r_idx},{c_idx}]: {text[:60]}")

detailed_analysis('2025년도 문제.docx')
