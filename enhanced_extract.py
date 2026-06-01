from docx import Document
import json
import re
from collections import defaultdict

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

# 테이블에서 설명 추출
explanations_detail = {}
table_idx = 0
for idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text.startswith("설명:"):
                table_idx += 1
                lines_text = text.split('\n')
                
                # 첫 번째 라인에서 단원명 추출
                first_line = lines_text[0]
                category = ""
                if "단원명-" in first_line:
                    category = first_line.split("단원명-")[1].strip()
                
                # 설명 추출
                explanation_lines = []
                for line in lines_text[1:]:
                    cleaned = line.strip()
                    if cleaned:
                        if cleaned.startswith("-"):
                            cleaned = cleaned[1:].strip()
                        explanation_lines.append(cleaned)
                
                explanation = "\n".join(explanation_lines)
                explanations_detail[table_idx] = {
                    "category": category,
                    "explanation": explanation
                }

# 문단에서 문제 추출
lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

questions = []
i = 0
problem_idx = 1
current_title = "2025년도 제1회"

while i < len(lines):
    line = lines[i]
    
    # 회차 제목 감지
    if "년도" in line and "회" in line:
        current_title = line
        i += 1
        continue
    
    # 문제 찾기 (번호. 문제 형식)
    if re.match(r'^\d+\.', line):
        # 문제 문자열에서 정보 추출
        # 형식: "번호. 문제내용? 정답기호"
        parts = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])\s*$', line)
        
        if parts:
            question_num = int(parts.group(1))
            question_text = parts.group(2).strip()
            answer_char = parts.group(3)
            answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
            
            # 선택지 찾기
            i += 1
            options = []
            
            # 한 라인에 모든 선택지가 있는 경우
            if i < len(lines) and re.search(r'①.*②.*③.*④', lines[i]):
                option_line = lines[i]
                # 선택지 분리 (① ② ③ ④로 구분)
                option_line = re.sub(r'④\s+', '④ ', option_line)
                option_line = re.sub(r'③\s+', '③ ', option_line)
                option_line = re.sub(r'②\s+', '② ', option_line)
                option_line = re.sub(r'①\s+', '① ', option_line)
                
                parts_options = re.split(r'(?=[②③④])', option_line)
                
                for part in parts_options:
                    part = part.strip()
                    if part and part[0] in '①②③④':
                        options.append(part[1:].strip())
                
                i += 1
            
            # 선택지가 여러 줄에 걸쳐 있는 경우
            else:
                while i < len(lines):
                    option_line = lines[i].strip()
                    if re.match(r'^[①②③④]\s+', option_line):
                        match_opt = re.match(r'^[①②③④]\s+(.+)', option_line)
                        if match_opt:
                            options.append(match_opt.group(1))
                            i += 1
                        else:
                            break
                    elif re.match(r'^[①②③④]', option_line):
                        # 기호 직후에 공백 없는 경우
                        options.append(option_line[1:].strip())
                        i += 1
                    else:
                        break
            
            if len(options) == 4:
                # 설명과 단원명 가져오기
                exp_data = explanations_detail.get(problem_idx, {"category": "", "explanation": ""})
                
                questions.append({
                    'number': question_num,
                    'question': question_text,
                    'options': options,
                    'answer': answer_idx,
                    'category': exp_data['category'],
                    'explanation': exp_data['explanation'],
                    'exam': current_title
                })
                problem_idx += 1
            else:
                i += 1
        else:
            i += 1
    else:
        i += 1

# JSON으로 저장
output = {
    'total': len(questions),
    'updated_date': '2025-06-01',
    'questions': questions
}

with open('questions_2025_full.json', 'w', encoding='utf-8') as f:
    json.dump(output, f, ensure_ascii=False, indent=2)

print(f"✓ 총 {len(questions)}개 문제 추출 및 분석 완료\n")
print("=== 추출된 문제 요약 ===\n")

# 단원별 통계
categories = defaultdict(int)
for q in questions:
    categories[q['category']] += 1

for cat, count in sorted(categories.items(), key=lambda x: -x[1]):
    print(f"- {cat}: {count}개")

print("\n=== 설명이 부족한 문제 ===\n")
for q in questions:
    if not q['explanation'] or len(q['explanation']) < 20:
        print(f"문제 {q['number']}: {q['question'][:50]}")
        print(f"  현재 설명: {q['explanation'][:60] if q['explanation'] else '없음'}\n")
