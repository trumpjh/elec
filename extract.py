"""
2025년도 전기기능사 문제 및 설명 추출 스크립트
문제.docx와 설명.docx에서 데이터를 추출하여 JSON으로 저장합니다.
"""

from docx import Document
import re
import json
from pathlib import Path
from collections import defaultdict

class QuestionExtractor:
    """문제 추출 클래스"""
    
    def __init__(self):
        self.questions = []  # 리스트 사용 (중복 허용)
        self.current_exam = None
        self.seen_problems = set()  # (exam, number) 중복 확인용
    
    def extract_from_docx(self, filepath):
        """docx 파일에서 문제 추출"""
        print(f"\n📖 문제 추출 중: {filepath}")
        doc = Document(filepath)
        
        # 빈 문단 제외 (인덱스 매칭을 위해)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        for i, text in enumerate(paragraphs):
            # 회차 감지
            if '제' in text and '회' in text:
                match = re.search(r'제(\d+)회', text)
                if match:
                    self.current_exam = f'제{match.group(1)}회'
                    print(f"   → {self.current_exam} 시작")
                continue
            
            # 문제 찾기: "숫자. " 형식
            match = re.match(r'^(\d+)\.\s+(.+)$', text)
            if match and self.current_exam:
                problem_num = int(match.group(1))
                problem_text_with_answer = match.group(2)
                
                # 정답 심볼 추출 (①②③④ 중 하나)
                answer_symbol = None
                answer_index = None
                problem_text = problem_text_with_answer
                
                # 마지막에 있는 심볼 찾기
                for symbol, idx in {'①': 0, '②': 1, '③': 2, '④': 3}.items():
                    if problem_text_with_answer.endswith(symbol):
                        answer_symbol = symbol
                        answer_index = idx
                        problem_text = problem_text_with_answer[:-len(symbol)].strip()
                        break
                
                if answer_symbol is None:
                    print(f"   ⚠️  경고: 문제 {problem_num}의 정답 심볼을 찾을 수 없음")
                    continue
                
                # 선택지 추출 (다음 문단에서)
                options = []
                if i + 1 < len(paragraphs):
                    next_para = paragraphs[i + 1]
                    
                    # 정규식으로 선택지 추출
                    option_pattern = r'[①②③④]\s*([^①②③④]+?)(?=[①②③④]|$)'
                    matches = re.findall(option_pattern, next_para)
                    options = [m.strip() for m in matches if m.strip()]
                
                if len(options) != 4:
                    print(f"   ⚠️  경고: 문제 {problem_num}의 선택지가 {len(options)}개 (4개 필요)")
                
                # 중복 확인 (같은 회차의 같은 문제)
                problem_key = (self.current_exam, problem_num)
                if problem_key in self.seen_problems:
                    print(f"   ℹ️  {self.current_exam} 문제 {problem_num} (중복)")
                else:
                    self.seen_problems.add(problem_key)
                
                # 문제 저장
                question = {
                    'number': problem_num,
                    'question': problem_text,
                    'options': options,
                    'answer': answer_index,
                    'exam': self.current_exam,
                    'explanation': ''  # 나중에 추가됨
                }
                
                self.questions.append(question)
        
        print(f"   ✓ {len(self.questions)}개 문제 추출 완료")
        return self.questions
    
    def save_json(self, output_file):
        """JSON으로 저장"""
        data = {
            'total': len(self.questions),
            'questions': self.questions
        }
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        print(f"   ✓ {output_file}에 저장됨 ({len(self.questions)}개 문제)")
        return output_file


class ExplanationExtractor:
    """설명 추출 클래스"""
    
    def __init__(self):
        self.explanations = {}  # {문제번호: 설명}
    
    def extract_from_docx(self, filepath):
        """docx 파일에서 설명 추출 (표 기반)"""
        print(f"\n📖 설명 추출 중: {filepath}")
        doc = Document(filepath)
        
        # 표의 개수 확인
        tables = doc.tables
        print(f"   → 총 {len(tables)}개 표 발견")
        
        # 각 표에서 설명 추출
        for table_idx, table in enumerate(tables):
            if table.rows and table.rows[0].cells:
                cell_text = table.rows[0].cells[0].text.strip()
                
                if cell_text:
                    # "설명: 단원명-전기설비" 부분 제거
                    explanation = cell_text
                    if '설명:' in explanation:
                        # "설명: " 이후의 내용 추출
                        parts = explanation.split('설명:')
                        if len(parts) > 1:
                            explanation = parts[1].strip()
                            # "단원명-" 제거
                            if '-' in explanation:
                                explanation = '-'.join(explanation.split('-')[1:]).strip()
                    
                    self.explanations[table_idx] = explanation
        
        print(f"   ✓ {len(self.explanations)}개 설명 추출 완료")
        return self.explanations
    
    def save_json(self, output_file):
        """JSON으로 저장"""
        examples = [
            {'table_index': idx, 'explanation': exp}
            for idx, exp in self.explanations.items()
        ]
        
        data = {
            'total': len(examples),
            'examples': examples
        }
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        print(f"   ✓ {output_file}에 저장됨 ({len(examples)}개 설명)")
        return output_file


class DataIntegrator:
    """문제와 설명 통합 클래스"""
    
    @staticmethod
    def integrate(questions, explanations):
        """문제와 설명 통합"""
        print(f"\n🔗 데이터 통합 중...")
        
        # 문제번호별 설명 맵 생성
        explanation_map = {}
        for table_idx, explanation in explanations.items():
            # 표의 인덱스가 문제 순서와 일치
            explanation_map[table_idx] = explanation
        
        # 통합
        integrated = []
        matched_count = 0
        
        for q in questions:
            # 설명 찾기 (문제번호 기반)
            for table_idx, explanation in explanation_map.items():
                # 간단한 매칭: 표 인덱스가 문제 순서와 일치한다고 가정
                # 더 정확한 매칭이 필요하면 문제 텍스트로 검색
                pass
            
            integrated.append(q)
        
        print(f"   ✓ {len(integrated)}개 문제 통합 완료")
        return integrated
    
    @staticmethod
    def save_json(questions, output_file):
        """JSON으로 저장"""
        data = {
            'total': len(questions),
            'questions': questions
        }
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        print(f"   ✓ {output_file}에 저장됨 ({len(questions)}개 문제)")
        return output_file


def analyze_and_match():
    """문제와 설명을 분석하여 매칭"""
    print("\n" + "="*60)
    print("📊 문제와 설명 매칭 및 통합")
    print("="*60)
    
    # 문제 로드
    with open('questions.json', 'r', encoding='utf-8') as f:
        q_data = json.load(f)
        questions = q_data['questions']
    
    # 설명 로드
    with open('example.json', 'r', encoding='utf-8') as f:
        e_data = json.load(f)
        examples = e_data['examples']
    
    print(f"\n문제 수: {len(questions)}")
    print(f"설명 수: {len(examples)}")
    
    # 설명 매칭: 표 인덱스로 설명을 문제에 추가
    # 표의 순서 = 문제의 순서 라고 가정
    for idx, q in enumerate(questions):
        if idx < len(examples):
            q['explanation'] = examples[idx].get('explanation', '')
    
    # 회차별 정보
    by_exam = defaultdict(list)
    for q in questions:
        by_exam[q['exam']].append(q)
    
    print(f"\n회차별 문제 수:")
    for exam in sorted(by_exam.keys()):
        count = len(by_exam[exam])
        print(f"  {exam}: {count}개")
    
    print(f"\nTotal: {sum(len(v) for v in by_exam.values())}개")
    
    # 통합된 데이터를 questions.json에 다시 저장
    integrated_data = {
        'total': len(questions),
        'questions': questions
    }
    
    with open('questions.json', 'w', encoding='utf-8') as f:
        json.dump(integrated_data, f, ensure_ascii=False, indent=2)
    
    print(f"\n✓ questions.json 업데이트됨 (설명 포함)")
    
    # 설명 추가 현황
    with_explanation = sum(1 for q in questions if q.get('explanation'))
    print(f"설명 추가됨: {with_explanation}개 / {len(questions)}개")


# ============================================================
# 메인 실행
# ============================================================

if __name__ == '__main__':
    print("\n" + "="*60)
    print("🚀 전기기능사 문제 및 설명 추출 시작")
    print("="*60)
    
    try:
        # 1단계: 문제 추출
        q_extractor = QuestionExtractor()
        questions = q_extractor.extract_from_docx('2025년도 문제.docx')
        q_extractor.save_json('questions.json')
        
        # 2단계: 설명 추출
        e_extractor = ExplanationExtractor()
        explanations = e_extractor.extract_from_docx('2025년도 설명.docx')
        e_extractor.save_json('example.json')
        
        # 3단계: 분석 및 매칭
        analyze_and_match()
        
        print("\n" + "="*60)
        print("✅ 모든 작업 완료!")
        print("="*60)
        print("\n생성된 파일:")
        print("  - questions.json: 모든 문제 데이터")
        print("  - example.json: 모든 설명 데이터")
        print("\n다음 단계: web 인터페이스 개발")
        
    except Exception as e:
        print(f"\n❌ 오류 발생: {e}")
        import traceback
        traceback.print_exc()
