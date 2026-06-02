"""
2025년도 문제.docx와 2025년도 설명.docx 파일을 읽어서 
JSON 파일로 변환하는 스크립트
"""

from docx import Document
import json
import re

def extract_questions_from_docx(docx_path):
    """2025년도 문제.docx에서 모든 문제와 답 추출"""
    doc = Document(docx_path)
    questions_list = []  # 딕셔너리 대신 리스트 사용
    current_exam = None
    seen = set()  # 회차별 문제 번호 추적
    
    # 모든 문단을 순회하면서 문제 추출
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 회차 찾기
        if '제' in text and '회' in text and '문제' in text:
            match = re.search(r'제(\d+)회', text)
            if match:
                current_exam = f"제{match.group(1)}회"
                seen.clear()  # 새 회차에서 초기화
                continue
        
        # 문제 번호 찾기
        match = re.match(r'^(\d+)\.\s+(.*)', text)
        if match:
            num = int(match.group(1))
            rest = match.group(2)
            
            # 같은 회차에서 중복 방지
            key = (current_exam, num)
            if key in seen:
                continue
            seen.add(key)
            
            # 정답 심볼 추출
            answer = None
            q_text = rest
            
            symbol_to_idx = {'①': 0, '②': 1, '③': 2, '④': 3}
            
            # 문장 끝에서부터 정답 심볼 찾기
            for symbol, idx in symbol_to_idx.items():
                if rest.endswith(symbol):
                    answer = idx
                    q_text = rest[:-len(symbol)].strip()
                    break
                elif ' ' + symbol in rest:
                    # 공백 뒤의 심볼 찾기
                    pos = rest.rfind(' ' + symbol)
                    after = rest[pos + len(' ' + symbol):].strip()
                    if len(after) == 0:  # 심볼이 문장 끝
                        answer = idx
                        q_text = rest[:pos].strip()
                        break
            
            # 다음 문단에서 선택지 추출
            options = []
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                
                # 선택지 추출
                option_pattern = r'[①②③④]\s*([^①②③④]+?)(?=[①②③④]|$)'
                matches = re.findall(option_pattern, next_text)
                if matches:
                    options = [m.strip() for m in matches if m.strip()]
            
            questions_list.append({
                'number': num,
                'question': q_text,
                'options': options[:4],
                'answer': answer,
                'exam': current_exam if current_exam else '미분류'
            })
    
    # 번호순으로 정렬
    questions_list.sort(key=lambda x: (x['exam'], x['number']))
    
    return {'total': len(questions_list), 'questions': questions_list}

def extract_explanations_from_docx(docx_path):
    """2025년도 설명.docx에서 설명 추출"""
    doc = Document(docx_path)
    explanations = {}
    
    # 표에서 설명 추출 (각 표는 하나의 설명)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                
                if '설명:' in cell_text:
                    # 문제 번호 찾기 (단원명-전기설비 전에 있는 번호)
                    # 형식: "설명: 단원명-전기설비\n- ..."
                    lines = cell_text.split('\n')
                    
                    # 첫 번째 라인에서 문제 번호 찾기
                    first_line = lines[0] if lines else ""
                    
                    # 설명 텍스트는 "-"로 시작하는 부분부터
                    explanation_lines = [line.strip() for line in lines[1:] if line.strip()]
                    explanation_text = '\n'.join(explanation_lines)
                    
                    # 문제 번호를 찾기 위해 이전 문단 확인
                    # (설명 표가 문제에 대응됨)
    
    # 다른 방식: 문단에서 문제 번호와 표의 순서 맞추기
    question_numbers = []
    for para in doc.paragraphs:
        text = para.text.strip()
        # 문제 번호 추출 (예: "1.", "11.", "2.")
        match = re.match(r'^(\d+)\.\s*$', text)
        if match:
            question_numbers.append(int(match.group(1)))
    
    # 표와 문제 번호 매칭
    table_idx = 0
    for q_num in question_numbers:
        if table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            explanation = ""
            
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    
                    # "설명:" 텍스트 제거 및 설명 추출
                    if '설명:' in cell_text:
                        # 불필요한 머리글 제거
                        explanation = re.sub(r'설명:\s*단원명[^-]*-\s*\w+\s*', '', cell_text).strip()
                        # "-" 기호로 시작하는 경우 유지
                        explanation = '\n'.join([line.strip() for line in explanation.split('\n') if line.strip()])
                    elif cell_text:
                        if explanation:
                            explanation += '\n' + cell_text
                        else:
                            explanation = cell_text
            
            if explanation:
                explanations[q_num] = {
                    'problem_number': q_num,
                    'category': '전기설비',
                    'explanation': explanation
                }
            
            table_idx += 1
    
    # 번호순으로 정렬
    sorted_explanations = [explanations[k] for k in sorted(explanations.keys()) if k in explanations]
    
    return {'total': len(sorted_explanations), 'examples': sorted_explanations}

def main():
    # 문제 추출
    print("문제.docx 파일을 읽고 있습니다...")
    questions_data = extract_questions_from_docx('2025년도 문제.docx')
    
    with open('questions.json', 'w', encoding='utf-8') as f:
        json.dump(questions_data, f, ensure_ascii=False, indent=2)
    print(f"✓ questions.json 생성 완료 ({len(questions_data['questions'])}개 문제)")
    
    # 설명 추출
    print("\n설명.docx 파일을 읽고 있습니다...")
    explanations_data = extract_explanations_from_docx('2025년도 설명.docx')
    
    with open('example.json', 'w', encoding='utf-8') as f:
        json.dump(explanations_data, f, ensure_ascii=False, indent=2)
    print(f"✓ example.json 생성 완료 ({len(explanations_data['examples'])}개 설명)")
    
    print("\n데이터 통합 중...")
    # 통합된 데이터 생성 (모든 75개 문제 포함)
    integrated_data = {
        'total': len(questions_data['questions']),
        'questions': []
    }
    
    # 설명 매핑 딕셔너리 생성
    explanation_map = {}
    for exp in explanations_data['examples']:
        explanation_map[exp['problem_number']] = exp['explanation']
    
    # 각 문제에 해당 설명 추가
    for q in questions_data['questions']:
        q_num = q['number']
        explanation = explanation_map.get(q_num, "")  # 없으면 빈 문자열
        
        integrated_data['questions'].append({
            'number': q['number'],
            'question': q['question'],
            'options': q['options'],
            'answer': q['answer'],
            'explanation': explanation,
            'exam': q['exam']
        })
    
    with open('integrated_questions.json', 'w', encoding='utf-8') as f:
        json.dump(integrated_data, f, ensure_ascii=False, indent=2)
    print(f"✓ integrated_questions.json 생성 완료 (모든 75개 문제 + 설명)")
    
    # 검증
    print("\n=== 검증 결과 ===")
    print(f"총 문제: {len(questions_data['questions'])}개")
    print(f"총 설명: {len(explanations_data['examples'])}개")
    print(f"설명이 있는 문제: {sum(1 for q in integrated_data['questions'] if q.get('explanation'))}")
    
    # 회차별 통계
    print("\n=== 회차별 문제 개수 ===")
    rounds = {}
    for q in integrated_data['questions']:
        exam = q.get('exam', '미분류')
        if exam not in rounds:
            rounds[exam] = 0
        rounds[exam] += 1
    
    for exam in sorted(rounds.keys()):
        print(f"{exam}: {rounds[exam]}개")

if __name__ == '__main__':
    main()
