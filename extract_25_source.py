"""
제1회 25번 원본 추출
"""
from docx import Document
import re

doc = Document('2025년도 문제.docx')

print("\n【제1회 25번 원본 추출】")
print("-" * 70)

current_exam = None
found = False

for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f'제{match.group(1)}회'
    
    # 문제 25 찾기
    if current_exam == '제1회' and text.startswith('25.'):
        print(f"\n문단 {para_idx}: {text[:80]}...")
        found = True
        
        # 다음 10개 문단 출력
        for i in range(1, 15):
            next_idx = para_idx + i
            if next_idx < len(doc.paragraphs):
                next_para = doc.paragraphs[next_idx]
                next_text = next_para.text.strip()
                
                # 다음 문제가 시작되면 멈추기
                if re.match(r'^\d+\.\s+', next_text):
                    print(f"\n(다음 문제 시작)")
                    break
                
                if next_text:
                    print(f"+{i}: {next_text[:70]}...")
        break

if not found:
    print("문제 25를 찾을 수 없습니다")
