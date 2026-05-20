from docx import Document
import json
import re

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

questions = []
lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

# 설명과 단원명 추출
explanations = {}
for idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text.startswith("설명:"):
                lines_text = text.split('\n')
                
                # 첫 번째 라인에서 단원명 추출
                first_line = lines_text[0]
                category = ""
                if "단원명-" in first_line:
                    category = first_line.split("단원명-")[1].strip()
                
                # 설명 추출
                explanation_lines = []
                for line in lines_text[1:]:
                    if line.startswith("-"):
                        explanation_lines.append(line[1:].strip())
                    elif line.strip():
                        explanation_lines.append(line.strip())
                
                explanation = " ".join(explanation_lines)
                explanations[idx + 1] = {
                    "category": category,
                    "explanation": explanation
                }

# 문제 추출
i = 0
problem_idx = 1
while i < len(lines):
    line = lines[i]
    
    # 문제 찾기 (번호. 문제 형식)
    if re.match(r'^\d+\.', line):
        match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])', line)
        if match:
            question_num = int(match.group(1))
            question_text = match.group(2)
            answer_char = match.group(3)
            answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
            
            # 다음 라인에서 선택지 찾기
            options = []
            i += 1
            
            # 한 라인에 모든 선택지가 있는 경우
            if i < len(lines) and re.search(r'①.*②.*③.*④', lines[i]):
                option_line = lines[i]
                # 선택지 분리
                parts = re.split(r'(?=[②③④])', option_line)
                for part in parts:
                    if part.startswith('①'):
                        options.append(part[1:].strip())
                    elif part.startswith('②'):
                        options.append(part[1:].strip())
                    elif part.startswith('③'):
                        options.append(part[1:].strip())
                    elif part.startswith('④'):
                        options.append(part[1:].strip())
            
            # 선택지가 여러 줄에 걸쳐 있는 경우
            else:
                while i < len(lines):
                    option_line = lines[i]
                    if re.match(r'^[①②③④]', option_line):
                        match_opt = re.match(r'^([①②③④])\s+(.+)', option_line)
                        if match_opt:
                            options.append(match_opt.group(2))
                            i += 1
                    else:
                        break
                i -= 1
            
            if len(options) == 4:
                # 설명과 단원명 가져오기
                exp_data = explanations.get(problem_idx, {"category": "", "explanation": ""})
                
                questions.append({
                    'number': question_num,
                    'question': question_text,
                    'options': options,
                    'answer': answer_idx,
                    'category': exp_data['category'],
                    'explanation': exp_data['explanation']
                })
                problem_idx += 1
    
    i += 1

# JSON으로 저장
with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump({'questions': questions}, f, ensure_ascii=False, indent=2)

print(f"총 {len(questions)}개 문제 추출됨\n")
for q in questions:
    print(f"문제 {q['number']}: {q['question'][:50]}")
    print(f"  단원: {q['category']}")
    print(f"  설명: {q['explanation'][:60]}...\n")
