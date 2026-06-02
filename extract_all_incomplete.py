from docx import Document
import re

doc = Document('2025년도 문제.docx')

print("=" * 80)
print("모든 불완전한 문제의 선택지 추출")
print("=" * 80)

# 단락들을 순회하면서 모든 문제 찾기
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 문제 번호와 내용 추출 (1. ... ~ 60. ...)
    match = re.match(r'^(\d+)\.\s+(.+?)\s+([①②③④])\s*$', text)
    
    if match:
        problem_num = int(match.group(1))
        question = match.group(2)
        answer_char = match.group(3)
        
        # 불완전한 문제인지 확인
        incomplete_nums = {4, 32, 40, 44, 60, 2, 14, 25, 42, 51, 53, 55, 58, 59}
        
        if problem_num in incomplete_nums:
            print(f"\n【문제 {problem_num}】")
            print(f"질문: {question}")
            print(f"정답: {answer_char}")
            print(f"선택지:")
            
            # 다음 단락들에서 선택지 찾기 (①, ②, ③, ④로 시작)
            option_count = 0
            for j in range(1, 6):
                if i+j < len(doc.paragraphs):
                    next_text = doc.paragraphs[i+j].text.strip()
                    
                    # 선택지인지 확인 (①②③④ 포함)
                    if re.match(r'^[①②③④]', next_text):
                        print(f"  {next_text}")
                        option_count += 1
                    # 선택지가 아닌 다른 텍스트가 나오면 중단
                    elif option_count > 0 and next_text and '①' not in next_text:
                        break
                    # 빈 줄은 무시
                    elif not next_text:
                        continue
                    else:
                        break

print("\n" + "=" * 80)
