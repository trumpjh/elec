from docx import Document
import json
import re

doc = Document('2025년도 문제.docx')

# 추출될 데이터
extracted_questions = []
current_exam = None

para_list = doc.paragraphs

# 단락을 순회하면서 문제 찾기
i = 0
while i < len(para_list):
    para = para_list[i]
    text = para.text.strip()
    
    # 회차 감지
    if '제1회' in text and '문제' in text:
        current_exam = '제1회'
        i += 1
        continue
    elif '제2회' in text and '문제' in text:
        current_exam = '제2회'
        i += 1
        continue
    elif '제3회' in text and '문제' in text:
        current_exam = '제3회'
        i += 1
        continue
    elif '제4회' in text and '문제' in text:
        current_exam = '제4회'
        i += 1
        continue
    
    if not current_exam or not text:
        i += 1
        continue
    
    # 문제 라인 감지
    match = re.match(r'^(\d+)\.\s+(.+)\?\s*([①②③④])\s*$', text)
    if match:
        num = int(match.group(1))
        question_text = match.group(2) + '?'
        answer_symbol = match.group(3)
        
        answer_map = {'①': 0, '②': 1, '③': 2, '④': 3}
        answer = answer_map[answer_symbol]
        
        current_question = {
            'number': num,
            'question': question_text,
            'options': [],
            'answer': answer,
            'exam': current_exam
        }
        
        # 다음 단락들에서 보기 수집
        j = i + 1
        options_dict = {}
        
        # 최대 10개 단락까지 탐색
        for k in range(j, min(j + 10, len(para_list))):
            next_text = para_list[k].text.strip()
            
            if not next_text:
                continue
            
            # 한 줄에 모든 보기가 있는 경우
            if re.search(r'^①.*②.*③.*④', next_text):
                parts = re.split(r'①\s*', next_text)
                if len(parts) > 1:
                    remaining = parts[1]
                    opts = re.split(r'②\s*|③\s*|④\s*', remaining)
                    opts = [o.strip() for o in opts if o.strip()]
                    if len(opts) == 4:
                        current_question['options'] = opts
                        i = k  # 이 단락까지 처리됨
                        break
            
            # 개별 보기 라인
            elif re.match(r'^[①②③④]\s+', next_text):
                symbol = next_text[0]
                option_text = re.sub(r'^[①②③④]\s+', '', next_text)
                symbol_map = {'①': 0, '②': 1, '③': 2, '④': 3}
                idx = symbol_map[symbol]
                options_dict[idx] = option_text
                
                # 모든 4개 보기를 수집했으면 종료
                if len(options_dict) == 4:
                    current_question['options'] = [options_dict[j] for j in range(4)]
                    i = k
                    break
            
            # 다음 문제를 만나면 중단
            elif re.match(r'^(\d+)\.\s+(.+)\?\s*([①②③④])\s*$', next_text):
                break
        
        extracted_questions.append(current_question)
    
    i += 1

# 테이블에서 설명 추출
explanations_list = []
for table in doc.tables:
    for row in table.rows:
        if len(row.cells) > 0:
            cell_text = row.cells[0].text.strip()
            
            if cell_text.startswith('설명:'):
                explanation_full = cell_text.replace('설명:', '', 1).strip()
                
                # 단원명 추출
                category = '미분류'
                explanation = explanation_full
                
                for cat in ['전기설비', '전기기기', '전기이론']:
                    if f'단원명-{cat}' in explanation_full:
                        category = cat
                        explanation = explanation_full.split(f'단원명-{cat}', 1)[1].strip()
                        break
                
                explanations_list.append({
                    'category': category,
                    'explanation': explanation
                })

# 최종 데이터 조합
print("=" * 80)
print("추출 결과:")
print("=" * 80)
print(f"총 문제: {len(extracted_questions)}")
print(f"총 설명: {len(explanations_list)}")

# 옵션 확인
no_options = sum(1 for q in extracted_questions if len(q['options']) == 0)
with_options = sum(1 for q in extracted_questions if len(q['options']) > 0)
print(f"보기 있음: {with_options}, 보기 없음: {no_options}")

# 샘플 출력
print("\n" + "=" * 80)
print("샘플 (처음 10개):")
print("=" * 80)
for i, q in enumerate(extracted_questions[:10]):
    has_opt = "✓" if q['options'] else "✗"
    print(f"{i+1}. {q['exam']} - 문제 {q['number']}: {has_opt}")
    if q['options']:
        print(f"   {q['options']}")

# 보기 없는 문제 리스트
print("\n" + "=" * 80)
print("보기가 없는 문제:")
print("=" * 80)
no_opt_problems = [q for q in extracted_questions if len(q['options']) == 0]
for q in no_opt_problems:
    print(f"  {q['exam']} - 문제 {q['number']}")

# JSON 생성
questions_output = {
    'total': len(extracted_questions),
    'questions': extracted_questions
}

explanations_output = {
    'total': len(explanations_list),
    'examples': [
        {
            'problem_number': extracted_questions[i]['number'],
            'category': explanations_list[i]['category'],
            'explanation': explanations_list[i]['explanation']
        }
        for i in range(min(len(extracted_questions), len(explanations_list)))
    ]
}

# 저장
with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(questions_output, f, ensure_ascii=False, indent=2)

with open('example.json', 'w', encoding='utf-8') as f:
    json.dump(explanations_output, f, ensure_ascii=False, indent=2)

print(f"\n✅ 저장됨:")
print(f"  questions.json: {len(extracted_questions)}개")
print(f"  example.json: {len(explanations_output['examples'])}개")
