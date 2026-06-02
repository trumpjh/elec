"""
이미지를 추출해서 저장하고 questions.json에 정보 추가
"""

from docx import Document
from docx.oxml import parse_xml
import re
import json
import os
from pathlib import Path

print("\n" + "="*70)
print("🖼️  이미지 추출 및 처리")
print("="*70)

# 1단계: 이미지 저장 폴더 생성
images_dir = 'images'
if not os.path.exists(images_dir):
    os.makedirs(images_dir)
    print(f"\n✓ {images_dir} 폴더 생성")
else:
    print(f"\n✓ {images_dir} 폴더 존재 확인")

# 2단계: docx에서 이미지 추출
doc = Document('2025년도 문제.docx')

print("\n【이미지 추출】")
print("-" * 70)

image_map = {}  # {rel_id: 저장된 파일명}
image_count = 0

for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        image_count += 1
        
        # 이미지 데이터 가져오기
        image_part = rel.target_part
        image_bytes = image_part.blob
        
        # 파일명 생성
        file_ext = rel.target_ref.split('.')[-1]
        filename = f"problem_image_{image_count}.{file_ext}"
        filepath = os.path.join(images_dir, filename)
        
        # 저장
        with open(filepath, 'wb') as f:
            f.write(image_bytes)
        
        image_map[rel.rId] = filename
        print(f"  이미지 {image_count}: {filename} 저장됨")

# 3단계: 이미지 포함 문제 식별
print("\n【이미지 포함 문제 식별】")
print("-" * 70)

problems_with_images = {}  # {(exam, number): image_file}
current_exam = None

for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 회차 감지
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f'제{match.group(1)}회'
    
    # 이미지 확인
    for run in para.runs:
        if run._element.drawing_lst:
            # 이 이미지가 어느 문제에 속하는지 찾기
            for i in range(para_idx - 1, -1, -1):
                prev_para = doc.paragraphs[i]
                prev_text = prev_para.text.strip()
                
                match = re.match(r'^(\d+)\.\s+(.+)$', prev_text)
                if match and prev_text:
                    problem_num = int(match.group(1))
                    
                    # 이미지 관계 확인
                    for rel_id, image_file in image_map.items():
                        if image_file:  # 매핑된 이미지 확인
                            problems_with_images[(current_exam, problem_num)] = image_file
                            print(f"  {current_exam} 문제 {problem_num} → {image_file}")
                    break

# 4단계: questions.json 업데이트
print("\n【questions.json 업데이트】")
print("-" * 70)

with open('questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 이미지 정보 추가
for question in data['questions']:
    exam = question['exam']
    number = question['number']
    
    if (exam, number) in problems_with_images:
        image_file = problems_with_images[(exam, number)]
        question['image'] = f"images/{image_file}"
        print(f"  {exam} 문제 {number}: 이미지 추가됨")

# 저장
with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print(f"\n✓ questions.json 업데이트 완료")

# 5단계: 결과 확인
print("\n【추가된 이미지 정보】")
print("-" * 70)

for question in data['questions']:
    if 'image' in question and question['image']:
        print(f"  {question['exam']} 문제 {question['number']}: {question['image']}")

print("\n" + "="*70)
print("✅ 모든 작업 완료!")
print("="*70)
