from docx import Document
import json
import re
from collections import defaultdict

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

# 모든 문제와 선택지 추출
questions_data = []
explanations_data = {}

# 1단계: 테이블에서 설명 추출
table_idx = 0
for idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text.startswith("설명:"):
                table_idx += 1
                lines_text = text.split('\n')
                
                first_line = lines_text[0]
                category = ""
                if "단원명-" in first_line:
                    category = first_line.split("단원명-")[1].strip()
                
                explanation_lines = []
                for line in lines_text[1:]:
                    cleaned = line.strip()
                    if cleaned:
                        if cleaned.startswith("-"):
                            cleaned = cleaned[1:].strip()
                        explanation_lines.append(cleaned)
                
                explanation = "\n".join(explanation_lines)
                explanations_data[table_idx] = {
                    "category": category,
                    "explanation": explanation
                }

# 2단계: 모든 단락을 한 줄의 텍스트로 만들기
paragraphs_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        paragraphs_text.append(text)

# 3단계: 문제 추출 (개선된 방식)
problem_idx = 1
i = 0
current_exam = "기출문제"

while i < len(paragraphs_text):
    line = paragraphs_text[i]
    
    # 회차 표시 감지
    exam_match = re.search(r'제(\d+)회', line)
    if exam_match:
        current_exam = f"제{exam_match.group(1)}회"
        i += 1
        continue
    
    # 패턴 1: "번호. 문제? ④" (정규)
    match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])\s*$', line)
    
    if match:
        question_num = int(match.group(1))
        question_text = match.group(2).strip()
        answer_char = match.group(3)
        answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
        
        # 다음 줄에서 선택지 찾기
        i += 1
        options = []
        
        # 한 라인에 모든 선택지가 있는 경우
        if i < len(paragraphs_text) and re.search(r'①.*②.*③.*④', paragraphs_text[i]):
            option_line = paragraphs_text[i]
            parts = re.split(r'(?=[②③④])', option_line)
            for part in parts:
                part = part.strip()
                if part and part[0] in '①②③④':
                    options.append(part[1:].strip())
            i += 1
        
        # 선택지가 여러 줄에 걸쳐 있는 경우
        else:
            while i < len(paragraphs_text):
                option_line = paragraphs_text[i].strip()
                if re.match(r'^[①②③④]', option_line):
                    match_opt = re.match(r'^[①②③④]\s+(.+)', option_line)
                    if match_opt:
                        options.append(match_opt.group(1))
                        i += 1
                    else:
                        break
                else:
                    break
        
        if len(options) == 4:
            exp_data = explanations_data.get(problem_idx, {"category": "", "explanation": ""})
            
            questions_data.append({
                'number': question_num,
                'question': question_text,
                'options': options,
                'answer': answer_idx,
                'exam': current_exam
            })
            
            explanations_data[problem_idx] = {
                'number': question_num,
                'category': exp_data['category'],
                'explanation': exp_data['explanation']
            }
            
            problem_idx += 1
        else:
            pass  # 선택지 4개 미만 (오류)
    else:
        i += 1

# 4단계: JSON 저장
questions_output = {
    'total': len(questions_data),
    'questions': questions_data
}

with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(questions_output, f, ensure_ascii=False, indent=2)

# 설명 저장
example_data = []
for idx in sorted(explanations_data.keys()):
    if explanations_data[idx]['explanation']:
        example_data.append({
            'problem_number': explanations_data[idx].get('number', idx),
            'category': explanations_data[idx]['category'],
            'explanation': explanations_data[idx]['explanation']
        })

example_output = {
    'total': len(example_data),
    'examples': example_data
}

with open('example.json', 'w', encoding='utf-8') as f:
    json.dump(example_output, f, ensure_ascii=False, indent=2)

# 5단계: 회차별 집계
by_exam = defaultdict(list)
for q in questions_data:
    by_exam[q['exam']].append(q['number'])

print("=" * 80)
print("✓ 파일 생성 완료!")
print("=" * 80)
print(f"\n총 문제: {len(questions_data)}개")
print(f"총 설명: {len(example_data)}개")

print("\n회차별 분석:")
for exam in sorted(by_exam.keys(), key=lambda x: int(x[1]) if re.search(r'\d', x) else 0):
    counts = len(by_exam[exam])
    print(f"  {exam}: {counts}개")

print("\n=== 샘플 ===")
print(f"\nquestions.json 첫 문제:")
print(json.dumps(questions_data[0], ensure_ascii=False, indent=2))

print(f"\nexample.json 첫 설명:")
if example_data:
    print(json.dumps(example_data[0], ensure_ascii=False, indent=2))
