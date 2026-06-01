from docx import Document
import re

doc = Document('2025년도 문제.docx')

# 4개 문제의 선택지 추출
problems_info = {
    (13, '제1회'): {'para': 22},
    (39, '제1회'): {'para': 56},
    (49, '제1회'): {'para': 89},
    (2, '제2회'): {'para': 117}
}

for (num, exam), info in problems_info.items():
    para_idx = info['para']
    print(f"\n{'='*70}")
    print(f"[{exam}] 문제 {num} (단락 {para_idx} 주변)")
    print(f"{'='*70}")
    
    # 문제 내용
    problem_line = doc.paragraphs[para_idx].text.strip()
    print(f"문제: {problem_line}\n")
    
    # 다음 10개 라인 출력 (선택지 찾기)
    options = []
    for j in range(1, 10):
        if para_idx + j < len(doc.paragraphs):
            next_text = doc.paragraphs[para_idx + j].text.strip()
            if next_text:
                print(f"  [{j}] {next_text}")
                # 선택지 패턴 찾기
                if re.match(r'^[①②③④]', next_text):
                    match_opt = re.match(r'^([①②③④])\s+(.+)', next_text)
                    if match_opt:
                        options.append(match_opt.group(2))
    
    if len(options) >= 4:
        print(f"\n선택지 ({len(options)}개):")
        for i, opt in enumerate(options[:4]):
            print(f"  {['①', '②', '③', '④'][i]}: {opt[:50]}...")
