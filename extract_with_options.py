from docx import Document
import json
import re

doc = Document('2025년도 문제.docx')

# 추출될 데이터
extracted_questions = []
current_exam = None
current_question = None

# 단락을 순회하면서 문제와 보기 추출
for para in doc.paragraphs:
    text = para.text.strip()
    
    if not text:
        continue
    
    # 회차 감지
    if '제1회' in text and '문제' in text:
        current_exam = '제1회'
        continue
    elif '제2회' in text and '문제' in text:
        current_exam = '제2회'
        continue
    elif '제3회' in text and '문제' in text:
        current_exam = '제3회'
        continue
    elif '제4회' in text and '문제' in text:
        current_exam = '제4회'
        continue
    
    if not current_exam:
        continue
    
    # 문제 라인 감지: "N. 질문? ①/②/③/④"
    match = re.match(r'^(\d+)\.\s+(.+)\?\s*([①②③④])\s*$', text)
    if match:
        # 이전 문제가 있으면 저장
        if current_question is not None:
            extracted_questions.append(current_question)
        
        num = int(match.group(1))
        question_text = match.group(2) + '?'
        answer_symbol = match.group(3)
        
        answer_map = {'①': 0, '②': 1, '③': 2, '④': 3}
        answer = answer_map[answer_symbol]
        
        current_question = {
            'number': num,
            'question': question_text,
            'options': [],
            'answer': answer,
            'exam': current_exam
        }
    
    # 보기 라인 감지: "① ... ② ... ③ ... ④ ..."
    elif current_question is not None and re.search(r'^①.*②.*③.*④', text):
        # 보기 파싱
        # 첫 번째 ①를 기준으로 분리
        parts = re.split(r'①\s*', text)
        if len(parts) > 1:
            remaining = parts[1]
            
            # ②③④로 분리
            opts = re.split(r'②\s*|③\s*|④\s*', remaining)
            
            # 불필요한 빈 항목 제거
            opts = [o.strip() for o in opts if o.strip()]
            
            if len(opts) == 4:
                current_question['options'] = opts

# 마지막 문제 저장
if current_question is not None:
    extracted_questions.append(current_question)

# 테이블에서 설명 추출
explanations_list = []
for table in doc.tables:
    for row in table.rows:
        if len(row.cells) > 0:
            cell_text = row.cells[0].text.strip()
            
            if cell_text.startswith('설명:'):
                explanation_full = cell_text.replace('설명:', '', 1).strip()
                
                # 단원명 추출
                category = '미분류'
                explanation = explanation_full
                
                for cat in ['전기설비', '전기기기', '전기이론']:
                    if f'단원명-{cat}' in explanation_full:
                        category = cat
                        explanation = explanation_full.split(f'단원명-{cat}', 1)[1].strip()
                        break
                
                explanations_list.append({
                    'category': category,
                    'explanation': explanation
                })

# 최종 데이터 조합
print("=" * 80)
print("추출 결과:")
print("=" * 80)
print(f"총 문제: {len(extracted_questions)}")
print(f"총 설명: {len(explanations_list)}")

# 옵션 확인
no_options = sum(1 for q in extracted_questions if len(q['options']) == 0)
with_options = sum(1 for q in extracted_questions if len(q['options']) > 0)
print(f"보기 있음: {with_options}, 보기 없음: {no_options}")

# 샘플 출력
print("\n" + "=" * 80)
print("샘플 (처음 5개):")
print("=" * 80)
for i, q in enumerate(extracted_questions[:5]):
    print(f"\n{i+1}. {q['exam']} - 문제 {q['number']}")
    print(f"   질문: {q['question'][:50]}")
    print(f"   보기: {q['options']}")
    print(f"   정답: {q['answer']}")

# JSON 생성
questions_output = {
    'total': len(extracted_questions),
    'questions': extracted_questions
}

explanations_output = {
    'total': len(explanations_list),
    'examples': [
        {
            'problem_number': extracted_questions[i]['number'],
            'category': explanations_list[i]['category'],
            'explanation': explanations_list[i]['explanation']
        }
        for i in range(min(len(extracted_questions), len(explanations_list)))
    ]
}

# 저장
with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(questions_output, f, ensure_ascii=False, indent=2)

with open('example.json', 'w', encoding='utf-8') as f:
    json.dump(explanations_output, f, ensure_ascii=False, indent=2)

print(f"\n✅ 저장됨:")
print(f"  questions.json: {len(extracted_questions)}개")
print(f"  example.json: {len(explanations_output['examples'])}개")
