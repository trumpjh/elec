from docx import Document

doc = Document('2025년도 설명.docx')

# 44번 찾기 (더 넓은 범위)
print("=== 제1회 44번 주변 내용 (문단) ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    # 40번부터 50번까지 출력
    if '40.' in text or '41.' in text or '42.' in text or '43.' in text or '44.' in text or '45.' in text or '46.' in text or '47.' in text or '48.' in text or '49.' in text or '50.' in text:
        print(f"{i}: {text}")
    # 더 긴 내용이 있으면 전부 출력
    if i >= 40 and i <= 50:
        if text:
            print(f"{i}: {text[:150]}")

# 59번 찾기
print("\n\n=== 제3회 59번 주변 내용 (문단) ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '55.' in text or '56.' in text or '57.' in text or '58.' in text or '59.' in text or '60.' in text:
        print(f"{i}: {text}")
    if i >= 115 and i <= 130:
        if text:
            print(f"{i}: {text[:150]}")

# 섹션 44 전체 텍스트 가져오기
print("\n\n=== 문단 43-45 전체 ===")
for i in range(43, 46):
    if i < len(doc.paragraphs):
        print(f"{i}: {doc.paragraphs[i].text}")

print("\n\n=== 문단 123-125 전체 ===")
for i in range(123, 126):
    if i < len(doc.paragraphs):
        print(f"{i}: {doc.paragraphs[i].text}")
