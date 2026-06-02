from docx import Document

doc = Document('2025년도 문제.docx')

# 제1회 44번 찾기
print("=== 제1회 44번 전후 내용 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '제1회' in text or '44' in text and text.startswith('44'):
        # 주변 내용 출력
        if i > 0:
            print(f"{i-1}: {doc.paragraphs[i-1].text.strip()[:80]}")
        print(f">>> {i}: {text[:100]}")
        if i < len(doc.paragraphs) - 1:
            print(f"{i+1}: {doc.paragraphs[i+1].text.strip()[:80]}")
        print()

# 또는 테이블에서 44 찾기
print("\n=== 테이블에서 검색 ===")
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if '44' in cell.text:
                print(f"테이블 {table_idx}, 행 {row_idx}, 셀 {cell_idx}:")
                print(cell.text[:300])
                print()
