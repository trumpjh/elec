from docx import Document

doc = Document('2025년도 설명.docx')

# 모든 문장 출력해서 확인
print("=== 문서 내용 검색 (44 포함) ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '44' in text:
        # 앞뒤 5개 문단 출력
        print(f"\n--- 문단 {i} 주변 (44 발견) ---")
        for j in range(max(0, i-5), min(len(doc.paragraphs), i+10)):
            p_text = doc.paragraphs[j].text.strip()
            if p_text:
                print(f"{j}: {p_text[:100]}")

# 테이블도 확인
print("\n\n=== 테이블에서 44 찾기 ===")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if '44' in cell.text:
                print(f"발견: {cell.text[:200]}")
