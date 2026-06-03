from docx import Document
import re

doc = Document('2025년도 설명.docx')

# 제1회 관련 모든 항목
print("=== 제1회 관련 항목 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '제1회' in text and len(text) < 150:
        print(f"[{i}] {text}")

print("\n=== 제2회 관련 항목 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '제2회' in text and len(text) < 150:
        print(f"[{i}] {text}")

print("\n=== 제4회 관련 항목 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '제4회' in text and len(text) < 150:
        print(f"[{i}] {text}")
