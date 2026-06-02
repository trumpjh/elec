"""
제4회 문제 9의 설명과 카테고리 추출
"""
from docx import Document
import re

print("\n" + "="*70)
print("제4회 문제 9의 설명 찾기")
print("="*70)

# 제4회는 14번 문제까지 원래 있었으므로, 
# 설명.docx에서 제4회 14번 설명까지는 있을 것입니다.
# 하지만 문제 9의 설명이 있는지 확인해야 합니다.

doc = Document('2025년도 설명.docx')

print(f"\n설명 문서 테이블 수: {len(doc.tables)}")

# 원래 74개 설명이 있었으므로, 제4회는 14번 문제까지만 설명이 있을 것입니다.
# 제4회 문제 9의 설명을 수동으로 찾아봅시다.

# 제4회의 테이블들:
# 제1회: 0-31 (32개)
# 제2회: 32-47 (16개)  
# 제3회: 48-59 (12개)
# 제4회: 60-73 (14개)

# 제4회 문제 9는 제4회의 0번째... 아니, 번호가 중복되므로 다시 생각해야 합니다.
# 제4회 내 문제들: 2, 3, 5, 6, 8, (이제 9), 10, 11, 12, 13, 14, ...
# 하지만 원래 데이터에는 9가 없었으므로, 설명도 없을 가능성이 높습니다.

# 일단 제4회 테이블들을 살펴봅시다.
print("\n【제4회 테이블들 확인 (60-73)】")
print("-" * 70)

for table_idx in range(60, min(74, len(doc.tables))):
    table = doc.tables[table_idx]
    if table.rows:
        first_cell = table.rows[0].cells[0]
        cell_text = first_cell.text.strip()[:60] if first_cell.text else "[빈 셀]"
        print(f"테이블 {table_idx}: {cell_text}...")

print("\n⚠️  주의: 원래 데이터에는 제4회 문제 9가 없었으므로,")
print("설명.docx에도 해당 설명이 없을 가능성이 높습니다.")
print("\n수동으로 설명을 작성하거나, 별도의 설명을 찾아야 합니다.")

# 일단 제4회 문제 9와 관련된 키워드로 검색해봅시다.
print("\n【관련 설명 검색】")
print("-" * 70)

search_keywords = ['트라이액', '위상제어', '전력제어']

for keyword in search_keywords:
    found = False
    for table_idx, table in enumerate(doc.tables):
        if table.rows:
            for row in table.rows:
                for cell in row.cells:
                    if keyword in cell.text:
                        print(f"✓ {keyword} 발견 (테이블 {table_idx})")
                        print(f"  내용: {cell.text[:60]}...")
                        found = True
                        break
                if found:
                    break
