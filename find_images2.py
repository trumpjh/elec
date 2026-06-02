"""
2025년도 문제.docx에서 이미지 찾기 - 다시 시도
"""

from docx import Document
import re

print("\n" + "="*70)
print("🖼️  이미지 포함 문제 분석")
print("="*70)

doc = Document('2025년도 문제.docx')

# 1단계: 모든 이미지 관계 추출
print("\n【문서의 모든 이미지】")
print("-" * 70)

image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        image_count += 1
        print(f"  이미지 {image_count}: {rel.target_ref}")

print(f"\n총 {image_count}개 이미지 발견")

# 2단계: 각 문단에서 이미지 참조 확인
print("\n【문단별 이미지 확인】")
print("-" * 70)

current_exam = None
problems_with_images = []

for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 회차 감지
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f'제{match.group(1)}회'
    
    # 문단의 이미지 확인
    has_image = False
    for run in para.runs:
        # drawing 요소 확인
        if run._element.drawing_lst:
            has_image = True
            break
    
    if has_image:
        print(f"\n📍 문단 {para_idx}: 이미지 포함")
        if text:
            print(f"   텍스트: {text[:60]}...")
        
        # 이전 문단에서 문제 찾기
        for i in range(para_idx - 1, -1, -1):
            prev_para = doc.paragraphs[i]
            prev_text = prev_para.text.strip()
            
            # 문제 번호 패턴 확인
            match = re.match(r'^(\d+)\.\s+(.+)$', prev_text)
            if match and prev_text:
                problem_num = int(match.group(1))
                problem_text = match.group(2)
                
                print(f"   → {current_exam} 문제 {problem_num}")
                print(f"   → 내용: {problem_text[:50]}...")
                
                problems_with_images.append({
                    'exam': current_exam,
                    'number': problem_num,
                    'text': problem_text,
                    'para_idx': para_idx
                })
                break

# 3단계: 표의 이미지 확인
print("\n【표의 이미지】")
print("-" * 70)

for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                has_image = any(run._element.drawing_lst for run in para.runs)
                if has_image:
                    print(f"표 {table_idx}, 행 {row_idx}, 열 {col_idx}: 이미지 포함")

# 결과 출력
print("\n【이미지 포함 문제 목록】")
print("-" * 70)

if problems_with_images:
    print(f"\n총 {len(problems_with_images)}개 문제에 이미지 포함:\n")
    for prob in problems_with_images:
        print(f"  {prob['exam']} 문제 {prob['number']}")
else:
    print("\n이미지가 포함된 문제가 없습니다.")
    print("(또는 이미지가 다른 형식으로 저장되었을 가능성)")

print("\n" + "="*70)
