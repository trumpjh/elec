from docx import Document

doc = Document('2025년도 문제.docx')

# 불완전한 문제들
incomplete_problems = {
    '제1회': [4, 32, 40, 44, 60],
    '제2회': [2, 14, 25],
    '제3회': [42, 51, 53, 55, 58, 59]
}

print("=" * 80)
print("DOCX 파일에서 불완전한 문제들의 선택지 추출")
print("=" * 80)

# 모든 단락 출력하여 문제 위치 파악
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 제1회 4번 찾기
    if text.startswith('4.') and '철심 재료' in text:
        print(f"\n단락 {i}: 제1회 4번")
        print(f"  {text}")
        
        # 다음 2줄 출력 (선택지)
        for j in range(1, 3):
            if i+j < len(doc.paragraphs):
                next_text = doc.paragraphs[i+j].text.strip()
                if next_text:
                    print(f"  +{j}: {next_text}")
    
    # 제1회 32번 찾기
    if text.startswith('32.') and '회전 방향' in text:
        print(f"\n단락 {i}: 제1회 32번")
        print(f"  {text}")
        
        for j in range(1, 3):
            if i+j < len(doc.paragraphs):
                next_text = doc.paragraphs[i+j].text.strip()
                if next_text:
                    print(f"  +{j}: {next_text}")
    
    # 제1회 40번 찾기
    if text.startswith('40.') and '기자력' in text:
        print(f"\n단락 {i}: 제1회 40번")
        print(f"  {text}")
        
        for j in range(1, 3):
            if i+j < len(doc.paragraphs):
                next_text = doc.paragraphs[i+j].text.strip()
                if next_text:
                    print(f"  +{j}: {next_text}")
    
    # 제1회 44번 찾기
    if text.startswith('44.') and '송전 방식' in text:
        print(f"\n단락 {i}: 제1회 44번")
        print(f"  {text}")
        
        for j in range(1, 3):
            if i+j < len(doc.paragraphs):
                next_text = doc.paragraphs[i+j].text.strip()
                if next_text:
                    print(f"  +{j}: {next_text}")
    
    # 제1회 60번 찾기
    if text.startswith('60.') and '특수 장소' in text:
        print(f"\n단락 {i}: 제1회 60번")
        print(f"  {text}")
        
        for j in range(1, 5):
            if i+j < len(doc.paragraphs):
                next_text = doc.paragraphs[i+j].text.strip()
                if next_text:
                    print(f"  +{j}: {next_text}")

print("\n(제2회, 제3회 문제도 유사하게 찾을 수 있습니다...)")
