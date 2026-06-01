from docx import Document
import re

doc = Document('2025년도 문제.docx')

# 4개의 누락된 문제 찾기: 13, 39, 49 (1회), 2 (2회)
target_problems = {13, 39, 49, 2}
found = {}

current_exam = ""

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if not text:
        continue
    
    # 회차 표시
    exam_match = re.search(r'제(\d+)회', text)
    if exam_match:
        current_exam = f"제{exam_match.group(1)}회"
        continue
    
    # 문제 번호 찾기
    for target in target_problems:
        if text.startswith(f"{target}."):
            print(f"\n{'='*70}")
            print(f"[{current_exam}] 문제 {target}")
            print(f"{'='*70}")
            print(f"단락 {i}: {text}")
            found[target] = {'exam': current_exam, 'para': i, 'text': text}

# 누락된 문제 정보 출력
print("\n" + "=" * 80)
print("📍 누락된 4개 문제 정보")
print("=" * 80)

# 정답 찾기 (다음 문단들 확인)
for target in sorted(target_problems):
    if target in found:
        para_idx = found[target]['para']
        print(f"\n문제 {target} (단락 {para_idx}):")
        print(f"  {found[target]['text'][:60]}...")
        
        # 다음 5개 문단 확인
        print(f"  이후 내용:")
        for j in range(1, 8):
            if para_idx + j < len(doc.paragraphs):
                next_text = doc.paragraphs[para_idx + j].text.strip()
                if next_text and len(next_text) > 0:
                    print(f"    [{j}] {next_text[:55]}...")
                    # 정답 찾기 (④ 등)
                    if any(ans in next_text for ans in ['①', '②', '③', '④']):
                        ans_match = re.search(r'([④③②①])', next_text)
                        if ans_match:
                            print(f"       → 정답 발견: {ans_match.group(1)}")
