"""
제1회 문제 4 찾기
"""
from docx import Document
import re

doc = Document('2025년도 문제.docx')

print("\n【제1회 문제 4 찾기】")
print("-" * 70)

current_exam = None
for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f'제{match.group(1)}회'
    
    # 문제 4 찾기
    if current_exam == '제1회' and text.startswith('4.'):
        print(f"\n문단 {para_idx}: {text[:70]}...")
        
        # 다음 10개 문단 출력
        for i in range(10):
            next_idx = para_idx + i
            if next_idx < len(doc.paragraphs):
                next_para = doc.paragraphs[next_idx]
                next_text = next_para.text.strip()
                marker = " [문제 시작]" if next_idx == para_idx else ""
                print(f"  +{i} (문단 {next_idx}): {next_text[:60]}...{marker}")
        break
