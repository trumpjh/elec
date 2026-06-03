from docx import Document
import re

doc = Document('2025년도 설명.docx')

# 모든 문단과 표를 순회
paragraphs = doc.paragraphs
tables = doc.tables

current_exam = None
for i, para in enumerate(paragraphs):
    text = para.text.strip()
    
    # 헤더에서 회차 추출
    if '제' in text and '회' in text and '설명' in text:
        match = re.search(r'제(\d)회', text)
        if match:
            current_exam = f"제{match.group(1)}회"
            print(f"[헤더] {text}")
    
    # 문제 번호 감지
    if re.match(r'^44\.$', text):
        print(f"\n✓ 찾음: {current_exam} 44번")
        # 다음 표 찾기
        for j in range(i, len(paragraphs)):
            # 이 방법은 작동하지 않을 수 있으므로 다른 방식으로 수정해야 함
        break

# 다른 방식: body element를 사용
current_exam = None
found_44 = False

for element in doc.element.body:
    if element.tag.endswith('p'):
        # 문단
        p = element
        text = ''.join([t.text for t in p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')])
        
        # 헤더 확인
        if '제' in text and '회' in text and '설명' in text:
            match = re.search(r'제(\d)회', text)
            if match:
                current_exam = f"제{match.group(1)}회"
        
        # 44번 확인
        if re.match(r'^44\.$', text.strip()):
            if current_exam == '제1회':
                print(f"\n✓ 찾음: {current_exam} 44번")
                found_44 = True
    
    elif element.tag.endswith('tbl') and found_44:
        # 표
        from docx.table import Table
        table = Table(element, doc.element.body)
        if len(table.rows) > 0:
            explanation_text = table.rows[0].cells[0].text
            print(f"설명 내용:\n{explanation_text}\n")
            found_44 = False
