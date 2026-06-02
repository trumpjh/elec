from docx import Document

doc = Document('2025년도 문제.docx')

print("=" * 70)
print("제4회 전체 내용 (단락 230부터)")
print("=" * 70)

# 제4회 시작 단락
start_idx = 230

# 제4회의 모든 내용 출력
for i in range(start_idx, min(start_idx + 50, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    if para.text.strip():  # 빈 줄 제외
        print(f"{i}: {para.text}")
    
    # "9" 또는 "9번"이 나오면 종료 (8번 이후)
    if '9번' in para.text or (para.text.startswith('9.') and len(para.text) > 3):
        break
