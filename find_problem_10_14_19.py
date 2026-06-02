from docx import Document

doc = Document('2025년도 문제.docx')

print("=" * 70)
print("제4회 전체 문제 내용")
print("=" * 70)

# 제4회 시작 단락 (230)
start_idx = 230

# 제4회의 모든 내용 출력
content_lines = []
for i in range(start_idx, len(doc.paragraphs)):
    para = doc.paragraphs[i]
    if para.text.strip():
        content_lines.append((i, para.text))
    
    # 다음 회차가 나오면 종료
    if i > start_idx and '제' in para.text and '회' in para.text and para.text.startswith('2025'):
        break

# 10번, 14번, 19번 찾기
for idx, text in content_lines:
    if text.startswith('10.') or text.startswith('14.') or text.startswith('19.'):
        print(f"\n단락 {idx}: {text}")
        
        # 다음 3줄 출력 (선택지)
        for i in range(1, 4):
            if (idx + i) < len(doc.paragraphs):
                next_text = doc.paragraphs[idx + i].text.strip()
                if next_text:
                    print(f"  +{i}: {next_text}")
                if '①' in next_text or '②' in next_text:
                    # 선택지가 시작됨, 계속 출력
                    for j in range(i+1, min(i+5, len(doc.paragraphs))):
                        next_next_text = doc.paragraphs[idx + j].text.strip()
                        if next_next_text and ('①' in next_next_text or '②' in next_next_text or '③' in next_next_text or '④' in next_next_text):
                            print(f"  +{j}: {next_next_text}")
                        else:
                            break
                    break
