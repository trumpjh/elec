"""
옵션 문제가 있는 문제들 재추출
"""
from docx import Document
import re
import json

print("\n" + "="*70)
print("옵션이 부족한 문제들 재추출")
print("="*70)

doc = Document('2025년도 문제.docx')

# 옵션이 1개인 문제들
problem_to_fix = [
    ('제1회', 4),
    ('제1회', 25),
    ('제1회', 32),
    ('제1회', 40),
    ('제1회', 44),
]

print("\n【문제별 문단 탐색】")
print("-" * 70)

current_exam = None
problem_paragraphs = {}  # {(exam, number): [문단들]}

for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 회차 감지
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f'제{match.group(1)}회'
    
    # 문제 번호 확인
    match = re.match(r'^(\d+)\.\s+(.+)$', text)
    if match:
        problem_num = int(match.group(1))
        key = (current_exam, problem_num)
        
        if key in dict(problem_to_fix):
            print(f"\n{current_exam} 문제 {problem_num} (문단 {para_idx}):")
            print(f"  {text[:70]}...")
            
            # 다음 4개 문단 확인 (선택지)
            options = []
            for i in range(1, 5):
                if para_idx + i < len(doc.paragraphs):
                    next_para = doc.paragraphs[para_idx + i]
                    next_text = next_para.text.strip()
                    
                    # 선택지 패턴 확인
                    opt_match = re.match(r'^[①②③④]\s*(.+)', next_text)
                    if opt_match:
                        options.append(opt_match.group(1))
                        print(f"    옵션 {len(options)}: {options[-1][:50]}...")
                    else:
                        print(f"    (문단 {para_idx + i}는 선택지 아님: {next_text[:30]}...)")
            
            if len(options) == 4:
                problem_paragraphs[key] = options

print("\n【재추출 완료】")
print("-" * 70)

for key, options in problem_paragraphs.items():
    exam, num = key
    print(f"{exam} 문제 {num}: {len(options)}개 옵션")

# questions.json 업데이트
print("\n【questions.json 업데이트】")
print("-" * 70)

with open('questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

updated_count = 0
for (exam, num), options in problem_paragraphs.items():
    for q in data['questions']:
        if q['exam'] == exam and q['number'] == num:
            if len(options) == 4:
                q['options'] = [f"{i+1}번: {options[i]}" if not options[i].startswith('①②③④'[i]) else options[i] for i in range(4)]
                # 더 깔끔하게 포맷
                q['options'] = options
                updated_count += 1
                print(f"✓ {exam} 문제 {num} 업데이트")
            break

print(f"\n총 {updated_count}개 문제 업데이트")

with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print("\n✓ questions.json 저장 완료")
