from docx import Document

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

print(f"Total paragraphs: {len(doc.paragraphs)}\n")
print("=== 모든 문단 ===\n")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        display = text[:120] if len(text) > 120 else text
        print(f"{i:3d}: {display}")
    else:
        print(f"{i:3d}: [empty]")
