from docx import Document
import json
import re

doc = Document('2025년도 문제.docx')

# 문제와 설명 정보 수집
questions_list = []
explanations_dict = {}

current_exam = None
table_idx = 0

# 1단계: 모든 문제 추출 (단락에서)
for para in doc.paragraphs:
    text = para.text.strip()
    
    # 회차 감지
    if '제1회' in text and '문제' in text:
        current_exam = '제1회'
        continue
    elif '제2회' in text and '문제' in text:
        current_exam = '제2회'
        continue
    elif '제3회' in text and '문제' in text:
        current_exam = '제3회'
        continue
    elif '제4회' in text and '문제' in text:
        current_exam = '제4회'
        continue
    
    if not current_exam or not text:
        continue
    
    # 문제 감지: "N. " 형식
    match = re.match(r'^(\d+)\.\s+(.+)\?\s*([①②③④])\s*$', text)
    if match:
        num = int(match.group(1))
        question_text = match.group(2) + '?'
        answer_symbol = match.group(3)
        
        # 답 매핑
        answer_map = {'①': 0, '②': 1, '③': 2, '④': 3}
        answer = answer_map[answer_symbol]
        
        q_obj = {
            'number': num,
            'question': question_text,
            'options': [],
            'answer': answer,
            'exam': current_exam
        }
        questions_list.append(q_obj)

# 2단계: 옵션 추가
options_lines = []
current_question_idx = None

for para in doc.paragraphs:
    text = para.text.strip()
    
    # 옵션 감지: "① ... ② ... ③ ... ④ ..."
    if '①' in text and '②' in text and '③' in text and '④' in text:
        # 옵션 파싱
        parts = re.split(r'[①②③④]\s+', text)
        parts = [p.strip() for p in parts if p.strip()]
        
        if len(parts) == 4 and current_question_idx is not None:
            questions_list[current_question_idx]['options'] = parts
    
    # 문제 번호 감지로 인덱스 업데이트
    match = re.match(r'^(\d+)\.\s+(.+)\?\s*([①②③④])\s*$', text)
    if match:
        current_question_idx = len(questions_list) - 1

# 3단계: 설명 추출 (테이블에서)
table_to_question_idx = 0

for table in doc.tables:
    if table_to_question_idx >= len(questions_list):
        break
    
    for row in table.rows:
        if len(row.cells) > 0:
            cell_text = row.cells[0].text.strip()
            
            if cell_text.startswith('설명:'):
                explanation_full = cell_text.replace('설명:', '', 1).strip()
                
                # 단원명 추출
                category = None
                explanation = explanation_full
                
                for cat in ['전기설비', '전기기기', '전기이론']:
                    if f'단원명-{cat}' in explanation_full:
                        category = cat
                        explanation = explanation_full.replace(f'단원명-{cat}', '').replace('-', '').strip()
                        break
                
                if table_to_question_idx < len(questions_list):
                    q_num = questions_list[table_to_question_idx]['number']
                    explanations_dict[q_num] = {
                        'category': category if category else '미분류',
                        'explanation': explanation
                    }
                    table_to_question_idx += 1

# 4단계: 분석
print("=" * 80)
print("추출 결과:")
print("=" * 80)
print(f"총 문제: {len(questions_list)}")
print(f"총 설명: {len(explanations_dict)}")

# 회차별 집계
exams = {}
for q in questions_list:
    exam = q['exam']
    if exam not in exams:
        exams[exam] = []
    exams[exam].append(q['number'])

for exam in sorted(exams.keys()):
    nums = exams[exam]
    unique_nums = len(set(nums))
    has_explanation = sum(1 for n in nums if n in explanations_dict)
    print(f"\n{exam}: {len(nums)}개 (고유 {unique_nums}개), 설명 {has_explanation}개")

# 누락된 설명
print("\n" + "=" * 80)
print("설명이 없는 문제:")
for q in questions_list:
    if q['number'] not in explanations_dict:
        print(f"  {q['exam']} - 문제 {q['number']}: {q['question'][:50]}")

print("\n" + "=" * 80)
print("상세: 첫 10개 문제")
for i in range(min(10, len(questions_list))):
    q = questions_list[i]
    has_exp = "✓" if q['number'] in explanations_dict else "✗"
    print(f"{i+1}. {q['exam']} - 문제 {q['number']}: {has_exp} {q['question'][:40]}")
