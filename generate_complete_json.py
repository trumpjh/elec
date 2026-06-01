from docx import Document
import json
import re
from collections import defaultdict

doc = Document(r"C:\Users\Administrator\Documents\history\elec\2025년도 문제.docx")

questions_data = []
explanations_data = {}

# 1. 설명 추출
table_idx = 0
for idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text.startswith("설명:"):
                table_idx += 1
                lines_text = text.split('\n')
                
                first_line = lines_text[0]
                category = ""
                if "단원명-" in first_line:
                    category = first_line.split("단원명-")[1].strip()
                
                explanation_lines = []
                for line in lines_text[1:]:
                    cleaned = line.strip()
                    if cleaned:
                        if cleaned.startswith("-"):
                            cleaned = cleaned[1:].strip()
                        explanation_lines.append(cleaned)
                
                explanation = "\n".join(explanation_lines)
                explanations_data[table_idx] = {
                    "category": category,
                    "explanation": explanation
                }

# 2. 문제 추출
paragraphs_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        paragraphs_text.append(text)

problem_idx = 1
i = 0
current_exam = "기출문제"

while i < len(paragraphs_text):
    line = paragraphs_text[i]
    
    exam_match = re.search(r'제(\d+)회', line)
    if exam_match:
        current_exam = f"제{exam_match.group(1)}회"
        i += 1
        continue
    
    match = re.match(r'^(\d+)\.\s+(.+)\?\s+([④③②①])\s*$', line)
    
    if match:
        question_num = int(match.group(1))
        question_text = match.group(2).strip()
        answer_char = match.group(3)
        answer_idx = {'①': 0, '②': 1, '③': 2, '④': 3}[answer_char]
        
        i += 1
        options = []
        
        if i < len(paragraphs_text) and re.search(r'①.*②.*③.*④', paragraphs_text[i]):
            option_line = paragraphs_text[i]
            parts = re.split(r'(?=[②③④])', option_line)
            for part in parts:
                part = part.strip()
                if part and part[0] in '①②③④':
                    options.append(part[1:].strip())
            i += 1
        else:
            while i < len(paragraphs_text):
                option_line = paragraphs_text[i].strip()
                if re.match(r'^[①②③④]', option_line):
                    match_opt = re.match(r'^[①②③④]\s+(.+)', option_line)
                    if match_opt:
                        options.append(match_opt.group(1))
                        i += 1
                    else:
                        break
                else:
                    break
        
        if len(options) == 4:
            exp_data = explanations_data.get(problem_idx, {"category": "", "explanation": ""})
            
            questions_data.append({
                'number': question_num,
                'question': question_text,
                'options': options,
                'answer': answer_idx,
                'exam': current_exam
            })
            
            explanations_data[problem_idx] = {
                'number': question_num,
                'category': exp_data['category'],
                'explanation': exp_data['explanation']
            }
            
            problem_idx += 1
    else:
        i += 1

# 3. 수동으로 누락된 4개 문제 추가
missing_problems = [
    {
        'number': 13,
        'question': '22.9[kV-Y] 가공전선의 굵기는 단면적이 몇 [mm²]이상이어야 하는가(단, 동선의 경우이다.)',
        'options': ['22', '32', '40', '50'],
        'answer': 0,
        'exam': '제1회'
    },
    {
        'number': 39,
        'question': '보극이 없는 직류기 운전 중 중성점의 위치가 변하지 않는 경우는',
        'options': ['과부하', '전부하', '중부하', '무부하'],
        'answer': 0,
        'exam': '제1회'
    },
    {
        'number': 49,
        'question': '전주 외등을 전주에 부착하는 경우 전주 외등은 하단으로부터 몇[m] 이상 높이에 시설하여야 하는가(단, 교통지장이 있는 경우이다.)',
        'options': ['3.0', '3.5', '4.0', '4.5'],
        'answer': 3,
        'exam': '제1회'
    },
    {
        'number': 2,
        'question': '그림과 같이 대전된 에보나이트 막대를 박검전기의 금속판에 닿지 않도록 가깝게 가져갔을 때 금박이 열렸다면 다음 중 옳은 것은(단, A는 원판, B는 박, C는 에보나이트 막대이다.)',
        'options': [
            'A: 양전기, B: 양전기, C: 음전기',
            'A: 음전기, B: 음전기, C: 음전기',
            'A: 양전기, B: 음전기, C: 음전기',
            'A: 양전기, B: 양전기, C: 양전기'
        ],
        'answer': 2,
        'exam': '제2회'
    }
]

# 누락된 문제 추가 (중복 제거)
for missing in missing_problems:
    if not any(q['number'] == missing['number'] and q['exam'] == missing['exam'] for q in questions_data):
        questions_data.append(missing)

# 문제 번호로 정렬
questions_data.sort(key=lambda x: (
    int(re.search(r'\d+', x['exam']).group()) if re.search(r'\d', x['exam']) else 0,
    x['number']
))

# 회차별 집계
by_exam = defaultdict(list)
for q in questions_data:
    by_exam[q['exam']].append(q['number'])

# JSON 저장
questions_output = {
    'total': len(questions_data),
    'questions': questions_data
}

with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(questions_output, f, ensure_ascii=False, indent=2)

# 설명 저장
example_data = []
for idx in sorted(explanations_data.keys()):
    if explanations_data[idx]['explanation']:
        example_data.append({
            'problem_number': explanations_data[idx].get('number', idx),
            'category': explanations_data[idx]['category'],
            'explanation': explanations_data[idx]['explanation']
        })

example_output = {
    'total': len(example_data),
    'examples': example_data
}

with open('example.json', 'w', encoding='utf-8') as f:
    json.dump(example_output, f, ensure_ascii=False, indent=2)

print("=" * 80)
print("✅ 파일 생성 완료!")
print("=" * 80)

print(f"\n📝 총 문제: {len(questions_data)}개")
print(f"📚 총 설명: {len(example_data)}개")

print("\n회차별 문제 개수:")
for exam in sorted(by_exam.keys(), key=lambda x: int(re.search(r'\d+', x).group()) if re.search(r'\d', x) else 0):
    counts = len(by_exam[exam])
    print(f"  {exam}: {counts}개")

print(f"\n예상 개수와 비교:")
expected = {"제1회": 32, "제2회": 16, "제3회": 12, "제4회": 15}
for exam, exp_count in expected.items():
    actual_count = len(by_exam.get(exam, []))
    status = "✓" if actual_count == exp_count else "✗"
    print(f"  {status} {exam}: {actual_count}개 (예상: {exp_count}개)")

print(f"\n✓ 샘플 문제:")
print(json.dumps(questions_data[0], ensure_ascii=False, indent=2))
