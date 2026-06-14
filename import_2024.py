"""
2024년도 문제/설명 docx를 추출하여 questions.json에 통합합니다.
"""

from __future__ import annotations

import json
import re
import shutil
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.table import Table

PROBLEM_DOCX = Path(r"d:\다운로드\2024년도 문제.docx")
EXPLANATION_DOCX = Path(r"d:\다운로드\2024년도 설명.docx")
QUESTIONS_JSON = Path("questions.json")
IMAGES_DIR = Path("images")
YEAR = 2024

# docx에 정답 심볼이 없는 문제 (exam_num, number) -> answer index
MANUAL_ANSWERS: dict[tuple[int, int], int] = {
    (1, 13): 2,  # ③ 온도퓨즈 (계기 내부 보호용)
}

SYMBOL_TO_INDEX = {"①": 0, "②": 1, "③": 2, "④": 3}
OPTION_PATTERN = r"[①②③④]\s*([^①②③④]+?)(?=[①②③④]|$)"


def strip_option_symbol(text: str) -> str:
    return re.sub(r"^[①②③④]\s*", "", text.strip())


def extract_category(explanation: str) -> str:
    if not explanation:
        return "기타"

    text = explanation
    if "설명:" in text:
        text = text.split("설명:", 1)[1].strip()
    if "-" in text.split("\n")[0]:
        first = text.split("\n")[0]
        if first.startswith("단원명-"):
            text = first.split("-", 1)[1] + "\n" + "\n".join(text.split("\n")[1:])

    first_line = text.split("\n")[0].strip()
    if first_line in ("전기기기", "전기설비", "전기이론"):
        return first_line

    keywords = {
        "전기기기": ["전기기기", "유도전동기", "발전기", "변압기", "다이오드"],
        "전기설비": ["전기설비", "가공", "금속관", "접지", "배전", "전선", "애자"],
        "전기이론": ["전기이론", "비오", "자기", "직류", "교류", "임피던스", "공진"],
    }
    for category, words in keywords.items():
        if any(word in explanation for word in words):
            return category
    return "기타"


def parse_options_from_text(text: str) -> list[str]:
    matches = re.findall(OPTION_PATTERN, text)
    options = [m.strip() for m in matches if m.strip()]
    if len(options) == 4:
        return options

    if "②" in text or "③" in text or "④" in text:
        first_option = re.split(r"[②③④]", text)[0].strip()
        rest = [m.strip() for m in re.findall(r"[②③④]\s*([^②③④]+?)(?=[②③④]|$)", text)]
        options = ([first_option] if first_option else []) + rest
        if len(options) == 4:
            return [strip_option_symbol(o) for o in options]

    return [strip_option_symbol(o) for o in options]


def collect_options(paragraphs: list[str], start_idx: int) -> tuple[list[str], int]:
    options: list[str] = []
    idx = start_idx

    while idx < len(paragraphs) and len(options) < 4:
        text = paragraphs[idx]
        if re.match(r"^\d+\.\s+", text):
            break
        if "2024년도" in text and "회" in text:
            break

        parsed = parse_options_from_text(text)
        if parsed:
            options.extend(parsed)
            idx += 1
            if len(options) >= 4:
                break
            continue

        symbol_hits = []
        for symbol, symbol_idx in SYMBOL_TO_INDEX.items():
            if text.startswith(symbol):
                symbol_hits.append((symbol_idx, strip_option_symbol(text)))

        if symbol_hits:
            for _, option_text in sorted(symbol_hits):
                if option_text:
                    options.append(option_text)
            idx += 1
            continue

        if options:
            break
        idx += 1

    return options[:4], idx


def extract_questions(filepath: Path) -> list[dict]:
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    questions: list[dict] = []
    current_exam = None
    i = 0

    while i < len(paragraphs):
        text = paragraphs[i]

        exam_match = re.search(r"제(\d+)회", text)
        if exam_match and "년도" in text:
            current_exam = f"제{exam_match.group(1)}회"
            i += 1
            continue

        problem_match = re.match(r"^(\d+)\.\s*(.+)$", text)
        if problem_match and current_exam:
            number = int(problem_match.group(1))
            body = problem_match.group(2).strip()

            answer_index = None
            question_text = body
            for symbol, idx in SYMBOL_TO_INDEX.items():
                if body.endswith(symbol):
                    answer_index = idx
                    question_text = body[: -len(symbol)].strip()
                    break

            options, next_idx = collect_options(paragraphs, i + 1)
            i = next_idx

            if answer_index is None:
                exam_num = int(current_exam.replace("제", "").replace("회", ""))
                manual = MANUAL_ANSWERS.get((exam_num, number))
                if manual is None:
                    print(f"  경고: {current_exam} {number}번 정답 심볼 없음 (건너뜀)")
                    continue
                answer_index = manual
                print(f"  참고: {current_exam} {number}번 수동 정답 적용 -> {answer_index}")
            if len(options) != 4:
                print(f"  경고: {current_exam} {number}번 선택지 {len(options)}개")

            questions.append(
                {
                    "number": number,
                    "question": question_text,
                    "options": options,
                    "answer": answer_index,
                    "exam": current_exam,
                    "explanation": "",
                    "category": "기타",
                    "year": YEAR,
                }
            )
            continue

        i += 1

    return questions


def extract_explanations(filepath: Path) -> dict[tuple[int, int], str]:
    doc = Document(filepath)
    explanations: dict[tuple[int, int], str] = {}
    current_exam = None
    last_problem_num = None
    last_problem_exam = None

    for element in doc.element.body:
        if element.tag.endswith("p"):
            text = "".join(
                t.text
                for t in element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                )
                if t.text
            )

            if f"{YEAR}년도" in text and "설명" in text:
                match = re.search(r"제(\d)회", text)
                if match:
                    current_exam = int(match.group(1))

            if text and re.match(r"^\d+\.$", text.strip()) and current_exam:
                last_problem_num = int(text.strip()[:-1])
                last_problem_exam = current_exam

        elif element.tag.endswith("tbl"):
            table = Table(element, doc)
            if not table.rows or not table.rows[0].cells:
                continue
            explanation_text = table.rows[0].cells[0].text.strip()
            if last_problem_exam and last_problem_num is not None:
                explanations[(last_problem_exam, last_problem_num)] = explanation_text
                last_problem_num = None
                last_problem_exam = None

    return explanations


def attach_explanations(questions: list[dict], explanations: dict[tuple[int, int], str]) -> None:
    for question in questions:
        exam_num = int(question["exam"].replace("제", "").replace("회", ""))
        key = (exam_num, question["number"])
        explanation = explanations.get(key, "")
        question["explanation"] = explanation
        question["category"] = extract_category(explanation)


def setup_problem_image(questions: list[dict], source_image: Path | None) -> None:
    IMAGES_DIR.mkdir(exist_ok=True)
    target_name = "2024년도 문제 제1회 3번.png"
    target_path = IMAGES_DIR / target_name

    if source_image and source_image.exists():
        shutil.copy2(source_image, target_path)
    elif not target_path.exists():
        return

    for question in questions:
        if question["exam"] == "제1회" and question["number"] == 3:
            question["image"] = f"images/{target_name}"


def merge_questions(new_questions: list[dict]) -> dict:
    if QUESTIONS_JSON.exists():
        with QUESTIONS_JSON.open("r", encoding="utf-8") as f:
            data = json.load(f)
        kept = [q for q in data.get("questions", []) if q.get("year") != YEAR]
    else:
        kept = []

    merged = kept + new_questions
    return {"total": len(merged), "questions": merged}


def print_summary(questions: list[dict]) -> None:
    by_exam = defaultdict(int)
    for q in questions:
        by_exam[q["exam"]] += 1

    print(f"\n2024년도 추출: {len(questions)}문제")
    for exam in sorted(by_exam):
        print(f"  {exam}: {by_exam[exam]}개")

    missing_explanation = [q for q in questions if not q.get("explanation")]
    if missing_explanation:
        print(f"  설명 없음: {len(missing_explanation)}개")
        for q in missing_explanation:
            print(f"    - {q['exam']} {q['number']}번")


def main() -> None:
    asset_image = Path(
        r"C:\Users\Jin\.cursor\projects\d-git-elec\assets"
        r"\c__Users_Jin_AppData_Roaming_Cursor_User_workspaceStorage_empty-window_images_"
        r"2024_______1__3_-01a9a001-a4be-4051-aba9-933c66519758.png"
    )

    print("2024년도 문제 추출 중...")
    questions = extract_questions(PROBLEM_DOCX)

    print("2024년도 설명 추출 중...")
    explanations = extract_explanations(EXPLANATION_DOCX)
    attach_explanations(questions, explanations)
    setup_problem_image(questions, asset_image)

    print_summary(questions)

    data = merge_questions(questions)
    with QUESTIONS_JSON.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"\nquestions.json 업데이트 완료 (총 {data['total']}문제)")


if __name__ == "__main__":
    main()
