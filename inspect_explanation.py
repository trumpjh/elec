from docx import Document

doc = Document('2025년도 설명.docx')

# 모든 문단 출력 (번호와 함께)
print("=== 전체 문서 내용 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:  # 빈 줄 제외
        print(f"{i}: {text[:120]}")

# 테이블 확인
print("\n\n=== 테이블 내용 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n테이블 {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()[:100]
            if text:
                print(f"  Row {r_idx}, Cell {c_idx}: {text}")
