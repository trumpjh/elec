from docx import Document
import json
import re

# 설명.docx 재분석
doc = Document('2025년도 설명.docx')

# 1. 모든 문단을 읽으면서 구조 파악
print("설명.docx 구조 분석:\n")

current_exam = None
problem_explanations = {}  # (exam, number) -> explanation

# 첫 번째 패스: 문제 번호 위치 찾기
problem_positions = {}  # 문제 번호 -> 문단 인덱스
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 회차 헤더 감지
    if '2025년도' in text and '설명' in text:
        match = re.search(r'제(\d)회', text)
        if match:
            current_exam = int(match.group(1))
            print(f"제{current_exam}회 설명 시작")
    
    # 문제 번호 감지
    if text and re.match(r'^\d+\.$', text) and current_exam:
        problem_num = int(text[:-1])  # "1." -> 1
        problem_positions[(current_exam, problem_num)] = i
        print(f"  발견: 제{current_exam}회 {problem_num}번 (문단 {i})")

print(f"\n총 발견된 문제: {len(problem_positions)}개")

# 2. 각 문제 번호 다음의 테이블에서 설명 추출
print("\n설명 추출:")
for (exam, num), para_idx in sorted(problem_positions.items()):
    # 이 문단 다음에 테이블이 있는지 확인
    # 테이블은 문단과 별도로 저장되므로, doc.tables에서 찾아야 함
    # 대신 다음 비어있지 않은 문단들을 추출
    
    # 간단한 방식: 다음 문단들 중 비어있지 않은 것을 설명으로 추출
    explanation_lines = []
    for j in range(para_idx + 1, min(para_idx + 10, len(doc.paragraphs))):
        para_text = doc.paragraphs[j].text.strip()
        
        # 다음 문제 번호에 도달하면 멈춤
        if re.match(r'^\d+\.$', para_text):
            break
        
        # 회차 헤더에 도달하면 멈춤
        if '2025년도' in para_text and '설명' in para_text:
            break
        
        if para_text:
            explanation_lines.append(para_text)
    
    if explanation_lines:
        explanation = '\n'.join(explanation_lines)
        problem_explanations[(exam, num)] = explanation
        if len(explanation) > 80:
            print(f"제{exam}회 {num}번: {explanation[:80]}...")
        else:
            print(f"제{exam}회 {num}번: {explanation}")

print(f"\n추출된 설명: {len(problem_explanations)}개")
