from docx import Document
import json
import re

doc = Document('2025년도 문제.docx')

# 추출될 데이터
extracted_questions = []
extracted_explanations = []

current_exam = None
last_question_num = None

# 1단계: 문제 추출
for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 회차 감지
    if '제1회' in text and '문제' in text:
        current_exam = '제1회'
        print(f"단락 {para_idx}: {current_exam} 시작")
        continue
    elif '제2회' in text and '문제' in text:
        current_exam = '제2회'
        print(f"단락 {para_idx}: {current_exam} 시작")
        continue
    elif '제3회' in text and '문제' in text:
        current_exam = '제3회'
        print(f"단락 {para_idx}: {current_exam} 시작")
        continue
    elif '제4회' in text and '문제' in text:
        current_exam = '제4회'
        print(f"단락 {para_idx}: {current_exam} 시작")
        continue
    
    if not current_exam or not text:
        continue
    
    # 문제 라인: "1. 질문? ①"
    match = re.match(r'^(\d+)\.\s+(.+)\?\s*([①②③④])\s*$', text)
    if match:
        num = int(match.group(1))
        question_text = match.group(2) + '?'
        answer_symbol = match.group(3)
        
        answer_map = {'①': 0, '②': 1, '③': 2, '④': 3}
        answer = answer_map[answer_symbol]
        
        q_obj = {
            'number': num,
            'question': question_text,
            'options': [],
            'answer': answer,
            'exam': current_exam
        }
        extracted_questions.append(q_obj)
        last_question_num = len(extracted_questions) - 1
        print(f"단락 {para_idx}: {current_exam} 문제 {num} 추출")

# 2단계: 옵션 추가
for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if not text:
        continue
    
    # 옵션 라인: "① ... ② ... ③ ... ④ ..."
    if re.search(r'^①[^②]*②[^③]*③[^④]*④', text):
        parts = re.split(r'[①②③④]\s+', text)
        parts = [p.strip() for p in parts if p.strip()]
        
        if len(parts) == 4 and last_question_num is not None:
            extracted_questions[last_question_num]['options'] = parts

# 3단계: 설명 추출 및 매핑
explanation_list = []
question_counter = 0

for table_idx, table in enumerate(doc.tables):
    for row in table.rows:
        if len(row.cells) > 0:
            cell_text = row.cells[0].text.strip()
            
            if cell_text.startswith('설명:'):
                explanation_full = cell_text.replace('설명:', '', 1).strip()
                
                # 단원명 추출
                category = '미분류'
                explanation = explanation_full
                
                for cat in ['전기설비', '전기기기', '전기이론']:
                    if f'단원명-{cat}' in explanation_full:
                        category = cat
                        explanation = explanation_full.split(f'단원명-{cat}', 1)[1].strip()
                        break
                
                explanation_list.append({
                    'category': category,
                    'explanation': explanation
                })

# 4단계: 매핑 (순서로 연결)
print(f"\n총 추출 문제: {len(extracted_questions)}")
print(f"총 추출 설명: {len(explanation_list)}")

# 각 문제에 설명 할당 (인덱스 기반)
questions_final = []
explanations_final = []

for i, q in enumerate(extracted_questions):
    q_copy = q.copy()
    questions_final.append(q_copy)
    
    # 설명이 있으면 할당
    if i < len(explanation_list):
        exp_obj = {
            'problem_number': q['number'],
            'category': explanation_list[i]['category'],
            'explanation': explanation_list[i]['explanation']
        }
        explanations_final.append(exp_obj)
        print(f"  문제 {q['number']} -> 설명 할당")
    else:
        print(f"  문제 {q['number']} -> 설명 없음 ❌")

# 5단계: 파일 생성
questions_output = {
    'total': len(questions_final),
    'questions': questions_final
}

explanations_output = {
    'total': len(explanations_final),
    'examples': explanations_final
}

# 저장
with open('questions_correct.json', 'w', encoding='utf-8') as f:
    json.dump(questions_output, f, ensure_ascii=False, indent=2)

with open('examples_correct.json', 'w', encoding='utf-8') as f:
    json.dump(explanations_output, f, ensure_ascii=False, indent=2)

print(f"\n✅ 저장됨:")
print(f"  questions_correct.json: {len(questions_final)}개")
print(f"  examples_correct.json: {len(explanations_final)}개")
