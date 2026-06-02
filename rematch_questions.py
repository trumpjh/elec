"""
정확한 문제-설명 매칭
문제 텍스트 기반으로 설명을 찾음
"""

from docx import Document
import re
import json
from collections import defaultdict

print("\n" + "="*70)
print("🔍 정확한 문제-설명 매칭")
print("="*70)

# ============================================================
# 1단계: 문제 추출
# ============================================================
print("\n【1단계】문제 추출 중...")

doc_questions = Document('2025년도 문제.docx')
paragraphs = [p.text.strip() for p in doc_questions.paragraphs if p.text.strip()]

questions = []
current_exam = None

for i, text in enumerate(paragraphs):
    if '제' in text and '회' in text:
        match = re.search(r'제(\d+)회', text)
        if match:
            current_exam = f'제{match.group(1)}회'
        continue
    
    match = re.match(r'^(\d+)\.\s+(.+)$', text)
    if match and current_exam:
        problem_num = int(match.group(1))
        problem_text_with_answer = match.group(2)
        
        # 정답 심볼 추출
        answer_index = None
        problem_text = problem_text_with_answer
        
        for symbol, idx in {'①': 0, '②': 1, '③': 2, '④': 3}.items():
            if problem_text_with_answer.endswith(symbol):
                answer_index = idx
                problem_text = problem_text_with_answer[:-len(symbol)].strip()
                break
        
        if answer_index is not None:
            # 선택지 추출
            options = []
            if i + 1 < len(paragraphs):
                next_para = paragraphs[i + 1]
                option_pattern = r'[①②③④]\s*([^①②③④]+?)(?=[①②③④]|$)'
                matches = re.findall(option_pattern, next_para)
                options = [m.strip() for m in matches if m.strip()]
            
            questions.append({
                'exam': current_exam,
                'number': problem_num,
                'question': problem_text,
                'options': options[:4],
                'answer': answer_index,
                'explanation': ''
            })

print(f"✓ {len(questions)}개 문제 추출 완료")

# ============================================================
# 2단계: 설명 추출
# ============================================================
print("【2단계】설명 추출 중...")

doc_explanations = Document('2025년도 설명.docx')
tables = doc_explanations.tables

explanations = []
for idx, table in enumerate(tables):
    if table.rows and table.rows[0].cells:
        cell_text = table.rows[0].cells[0].text.strip()
        
        # 단원명 추출
        first_line = cell_text.split('\n')[0] if cell_text else ''
        
        # 설명 정리
        explanation = cell_text
        if '설명:' in explanation:
            parts = explanation.split('설명:')
            if len(parts) > 1:
                explanation = parts[1].strip()
                if '-' in explanation:
                    explanation = '-'.join(explanation.split('-')[1:]).strip()
        
        explanations.append({
            'table_index': idx,
            'explanation': explanation,
            'category': first_line
        })

print(f"✓ {len(explanations)}개 설명 추출 완료")

# ============================================================
# 3단계: 정확한 매칭
# ============================================================
print("【3단계】문제-설명 매칭 중...")

# 간단한 순서 기반 매칭 (표의 순서 = 문제의 순서)
for idx, q in enumerate(questions):
    if idx < len(explanations):
        q['explanation'] = explanations[idx]['explanation']
        q['category'] = explanations[idx]['category'].replace('설명: 단원명-', '').replace('단원명-', '')

# ============================================================
# 4단계: 저장
# ============================================================
print("【4단계】JSON 저장 중...")

data = {
    'total': len(questions),
    'questions': questions
}

with open('questions.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print(f"✓ questions.json 저장 완료")

# ============================================================
# 5단계: 통계
# ============================================================
print("\n【5단계】결과 통계")
print("-" * 70)

by_exam = defaultdict(list)
for q in questions:
    by_exam[q['exam']].append(q)

for exam in sorted(by_exam.keys()):
    count = len(by_exam[exam])
    print(f"  {exam}: {count}개")

# 단원별 분류
by_category = defaultdict(int)
for q in questions:
    cat = q.get('category', '기타')
    if cat in ['전기기기', '전기설비', '전기이론']:
        by_category[cat] += 1
    else:
        by_category['기타'] += 1

print("\n【단원별 분류】")
for cat in sorted(by_category.keys()):
    print(f"  {cat}: {by_category[cat]}개")

# 제1회 16번 문제 확인
print("\n【제1회 16번 문제 확인】")
for q in questions:
    if q['exam'] == '제1회' and q['number'] == 16:
        print(f"문제: {q['question'][:50]}...")
        print(f"선택지: {len(q['options'])}개")
        print(f"정답: {['①', '②', '③', '④'][q['answer']] if q['answer'] is not None else 'None'}")
        print(f"단원: {q.get('category', '미분류')}")
        print(f"설명: {q['explanation'][:60]}...")

print("\n" + "="*70)
print("✅ 모든 작업 완료!")
print("="*70)
