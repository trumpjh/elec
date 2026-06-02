"""
extract_2025_data.py의 추출 함수를 직접 테스트
"""

from docx import Document
import re

def extract_questions_from_docx_debug(docx_path):
    """디버그 모드로 문제 추출"""
    doc = Document(docx_path)
    questions = {}
    current_exam = None
    
    # 모든 문단을 순회하면서 문제 추출
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 회차 찾기
        if '제' in text and '회' in text and '문제' in text:
            match = re.search(r'제(\d+)회', text)
            if match:
                current_exam = f"제{match.group(1)}회"
                continue
        
        # 문제 번호 찾기
        match = re.match(r'^(\d+)\.\s+(.*)', text)
        if match:
            num = int(match.group(1))
            rest = match.group(2)
            
            # 정답 심볼 추출
            answer = None
            q_text = rest
            
            symbol_to_idx = {'①': 0, '②': 1, '③': 2, '④': 3}
            
            # 문장 끝에서부터 정답 심볼 찾기
            for symbol, idx in symbol_to_idx.items():
                if rest.endswith(symbol):
                    answer = idx
                    q_text = rest[:-len(symbol)].strip()
                    break
                elif ' ' + symbol in rest:
                    # 공백 뒤의 심볼 찾기
                    pos = rest.rfind(' ' + symbol)
                    after = rest[pos + len(' ' + symbol):].strip()
                    if len(after) == 0:  # 심볼이 문장 끝
                        answer = idx
                        q_text = rest[:pos].strip()
                        break
            
            # 다음 문단에서 선택지 추출
            options = []
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                
                # 선택지 추출
                option_pattern = r'[①②③④]\s*([^①②③④]+?)(?=[①②③④]|$)'
                matches = re.findall(option_pattern, next_text)
                if matches:
                    options = [m.strip() for m in matches if m.strip()]
            
            if num not in questions:
                questions[num] = {
                    'number': num,
                    'question': q_text,
                    'options': options[:4],
                    'answer': answer,
                    'exam': current_exam if current_exam else '미분류'
                }
            else:
                print(f"중복 발견: 문제 {num}")
    
    # 번호순으로 정렬
    sorted_questions = [questions[k] for k in sorted(questions.keys())]
    
    return {'total': len(sorted_questions), 'questions': sorted_questions}

# 테스트
result = extract_questions_from_docx_debug('2025년도 문제.docx')
print(f"추출된 문제: {result['total']}개")
print(f"\n첫 5개 문제:")
for q in result['questions'][:5]:
    print(f"  {q['number']}: {q['exam']}")
